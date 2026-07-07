from __future__ import annotations

import tempfile
from pathlib import Path

import pytest


@pytest.fixture(scope="function")
def temp_dir():
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)


@pytest.fixture
def docx_path(temp_dir: Path) -> Path:
    import docx

    path = temp_dir / "test.docx"
    doc = docx.Document()
    doc.add_paragraph("Hello world.")
    doc.add_paragraph("Second paragraph with more content.")
    doc.save(str(path))
    return path


@pytest.fixture
def docx_empty_path(temp_dir: Path) -> Path:
    import docx

    path = temp_dir / "empty.docx"
    doc = docx.Document()
    doc.save(str(path))
    return path


@pytest.fixture
def docx_corrupt_path(temp_dir: Path) -> Path:
    path = temp_dir / "corrupt.docx"
    path.write_bytes(b"this is not a valid zip file")
    return path


@pytest.fixture
def odt_path(temp_dir: Path) -> Path:
    from odf.opendocument import OpenDocumentText
    from odf.text import P

    path = temp_dir / "test.odt"
    doc = OpenDocumentText()
    doc.text.addElement(P(text="ODT paragraph one."))
    doc.text.addElement(P(text="ODT paragraph two."))
    doc.save(str(path))
    return path


@pytest.fixture
def odt_empty_path(temp_dir: Path) -> Path:
    from odf.opendocument import OpenDocumentText

    path = temp_dir / "empty.odt"
    doc = OpenDocumentText()
    doc.save(str(path))
    return path


@pytest.fixture
def epub_path(temp_dir: Path) -> Path:
    from ebooklib import epub

    path = temp_dir / "test.epub"
    book = epub.EpubBook()
    book.set_title("Test Book")
    book.set_identifier("test123")
    book.set_language("en")

    chap1 = epub.EpubHtml(title="Chapter 1", file_name="chap1.xhtml", lang="en")
    chap1.content = "<h1>Chapter One</h1><p>This is chapter one content.</p>"
    book.add_item(chap1)

    chap2 = epub.EpubHtml(title="Chapter 2", file_name="chap2.xhtml", lang="en")
    chap2.content = "<h1>Chapter Two</h1><p>This is chapter two content.</p>"
    book.add_item(chap2)

    book.toc = (
        epub.Link("chap1.xhtml", "Chapter 1", "chap1"),
        epub.Link("chap2.xhtml", "Chapter 2", "chap2"),
    )
    book.spine = [chap1, chap2]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    epub.write_epub(str(path), book)
    return path


@pytest.fixture
def epub_empty_path(temp_dir: Path) -> Path:
    from ebooklib import epub

    path = temp_dir / "empty.epub"
    book = epub.EpubBook()
    book.set_title("Empty Book")
    book.set_identifier("empty123")
    book.set_language("en")
    book.spine = []
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(str(path), book)
    return path


@pytest.fixture
def pdf_empty_path(temp_dir: Path) -> Path:
    from pypdf import PdfWriter

    path = temp_dir / "empty.pdf"
    writer = PdfWriter()
    with open(str(path), "wb") as f:
        writer.write(f)
    return path


@pytest.fixture
def pdf_corrupt_path(temp_dir: Path) -> Path:
    path = temp_dir / "corrupt.pdf"
    path.write_bytes(b"this is not a valid pdf file")
    return path


@pytest.fixture
def docx_chapters_path(temp_dir: Path) -> Path:
    import docx

    path = temp_dir / "chapters.docx"
    doc = docx.Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the introduction text.")
    doc.add_paragraph("More intro content.")
    doc.add_heading("Methods", level=1)
    doc.add_paragraph("We used these methods.")
    doc.add_heading("A subsection", level=2)
    doc.add_paragraph("Some detail here.")
    doc.add_heading("Results", level=1)
    doc.add_paragraph("The results are in.")
    doc.save(str(path))
    return path


@pytest.fixture
def docx_nested_headings_path(temp_dir: Path) -> Path:
    import docx

    path = temp_dir / "nested.docx"
    doc = docx.Document()
    doc.add_heading("Part One", level=1)
    doc.add_paragraph("Part one content.")
    doc.add_heading("Section 1.1", level=2)
    doc.add_paragraph("Section 1.1 content.")
    doc.add_heading("Section 1.2", level=3)
    doc.add_paragraph("Section 1.2 content.")
    doc.add_heading("Part Two", level=1)
    doc.add_paragraph("Part two content.")
    doc.save(str(path))
    return path


@pytest.fixture
def docx_no_headings_path(temp_dir: Path) -> Path:
    import docx

    path = temp_dir / "no_headings.docx"
    doc = docx.Document()
    doc.add_paragraph("Just a paragraph.")
    doc.add_paragraph("Another paragraph.")
    doc.save(str(path))
    return path


@pytest.fixture
def odt_chapters_path(temp_dir: Path) -> Path:
    from odf.opendocument import OpenDocumentText
    from odf.text import H
    from odf.text import P

    path = temp_dir / "chapters.odt"
    doc = OpenDocumentText()
    doc.text.addElement(H(outlinelevel=1, text="Introduction"))
    doc.text.addElement(P(text="Introduction text goes here."))
    doc.text.addElement(P(text="More introduction."))
    doc.text.addElement(H(outlinelevel=1, text="Methods"))
    doc.text.addElement(P(text="Methods description."))
    doc.save(str(path))
    return path


@pytest.fixture
def epub_chapters_path(temp_dir: Path) -> Path:
    from ebooklib import epub

    path = temp_dir / "chapters.epub"
    book = epub.EpubBook()
    book.set_identifier("chapters123")
    book.set_language("en")

    chap1 = epub.EpubHtml(title="Chapter 1", file_name="chap1.xhtml", lang="en")
    chap1.content = "<h1>Introduction</h1><p>This is intro text.</p>"
    book.add_item(chap1)

    chap2 = epub.EpubHtml(title="Chapter 2", file_name="chap2.xhtml", lang="en")
    chap2.content = "<h2>Methods</h2><p>Methods text.</p><h3>Sub-method</h3><p>Sub-method text.</p>"
    book.add_item(chap2)

    book.toc = (epub.Link("chap1.xhtml", "Chapter 1", "chap1"),)
    book.spine = [chap1, chap2]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(str(path), book)
    return path

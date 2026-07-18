from __future__ import annotations

import json
import os
import queue
import re
import subprocess
import sys
import threading
import time
import warnings
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import Callable
from urllib.parse import quote
from urllib.request import urlopen

import imageio_ffmpeg
import numpy as np

os.environ.setdefault("FFMPEG_BINARY", imageio_ffmpeg.get_ffmpeg_exe())
warnings.filterwarnings(
    "ignore",
    message="Couldn't find ffmpeg or avconv - defaulting to ffmpeg, but may not work",
    category=RuntimeWarning,
)
from tkinter import END
from tkinter import BooleanVar
from tkinter import DoubleVar
from tkinter import IntVar
from tkinter import Menu
from tkinter import StringVar
from tkinter import Text
from tkinter import Tk
from tkinter import Toplevel
from tkinter import filedialog
from tkinter import messagebox
from tkinter import ttk

from pydub import AudioSegment

MODEL_NAME = "tts_models/multilingual/multi-dataset/xtts_v2"
DEFAULT_SPEAKER = "Ana Florence"
ENGINE_AUTO = "Auto"
ENGINE_PIPER = "Piper"
ENGINE_XTTS = "XTTS v2"
ENGINE_POCKET = "Pocket TTS"
POCKET_DEFAULT_VOICE = "alba"
POCKET_VOICE_DIR = Path.cwd() / "voices" / "pocket"
POCKET_LANG_MAP: dict[str, str] = {
    "en": "english",
    "fr": "french",
    "de": "german",
    "pt": "portuguese",
    "it": "italian",
    "es": "spanish",
}
# --- Language / engine capability registry ---
# Human-readable names, ordered as we want them to appear in the Language menu.
# Covers every language any engine speaks (XTTS is the widest — see below).
LANGUAGE_NAMES: dict[str, str] = {
    "en": "English",
    "hu": "Hungarian",
    "fr": "French",
    "de": "German",
    "pt": "Portuguese",
    "it": "Italian",
    "es": "Spanish",
    "pl": "Polish",
    "tr": "Turkish",
    "ru": "Russian",
    "nl": "Dutch",
    "cs": "Czech",
    "ar": "Arabic",
    "zh-cn": "Chinese",
    "ko": "Korean",
    "ja": "Japanese",
    "hi": "Hindi",
}

# Languages each neural engine can speak. Piper is intentionally absent here:
# its language coverage is derived at runtime from the installed voice models
# (see piper_languages()), because the user can download more via the wizard.
# XTTS_LANGUAGES is transcribed from the Coqui XTTS v2 model config
# (TTS.tts.configs.xtts_config.XttsConfig.languages) — the full 17-language set.
XTTS_LANGUAGES: list[str] = [
    "en", "es", "fr", "de", "it", "pt", "pl", "tr", "ru",
    "nl", "cs", "ar", "zh-cn", "hu", "ko", "ja", "hi",
]
POCKET_LANGUAGES: list[str] = list(POCKET_LANG_MAP.keys())

# Pocket TTS ships these named voices (kyutai-labs/pocket-tts). Hardcoded so the
# UI can offer a dropdown without triggering the heavy pocket_tts import.
POCKET_PREDEFINED_VOICES: list[str] = [
    "alba", "anna", "vera", "eve", "mary", "jane",
    "fantine", "cosette", "eponine", "azelma",
    "jean", "marius", "javert", "charles", "paul", "george", "michael",
    "bill_boerst", "peter_yearsley", "stuart_bell", "caro_davy",
    "giovanni", "lola", "juergen", "rafael", "estelle",
]
POCKET_DEFAULT_VOICE_FOR_LANG: dict[str, str] = {
    "en": "alba",
    "fr": "estelle",
    "de": "juergen",
    "pt": "rafael",
    "it": "giovanni",
    "es": "lola",
}

# One-line "what can this engine do" text, surfaced under the selectors so the
# user can make an informed choice rather than guessing.
ENGINE_SUMMARIES: dict[str, str] = {
    ENGINE_AUTO: (
        "Auto picks the engine for you: a reference clip clones the voice with "
        "XTTS, otherwise the local Piper voice is used."
    ),
    ENGINE_PIPER: (
        "Piper — fast, fully offline. Speaks whichever languages you have voices "
        "for; download more in the Voice Wizard. No voice cloning."
    ),
    ENGINE_XTTS: (
        "XTTS v2 — highest quality, 7 languages. Use a built-in speaker or clone "
        "any voice from a reference clip. Larger download, slower on CPU."
    ),
    ENGINE_POCKET: (
        "Pocket TTS — lightweight neural voices for 6 languages. Choose a built-in "
        "voice or clone one from a reference clip."
    ),
}
PIPER_VOICE_DIR = Path.cwd() / "voices" / "piper"
APP_SETTINGS_PATH = Path.cwd() / "settings.json"
PIPER_VOICES_JSON_URL = "https://huggingface.co/rhasspy/piper-voices/resolve/main/voices.json?download=true"
PIPER_VOICE_OPTIONS = {
    "Hungarian | Anna | medium": {
        "code": "hu_HU-anna-medium",
        "xtts_language": "hu",
    },
    "English US | Lessac | medium": {
        "code": "en_US-lessac-medium",
        "xtts_language": "en",
    },
    "English GB | Alan | medium": {
        "code": "en_GB-alan-medium",
        "xtts_language": "en",
    },
}
DEFAULT_PIPER_VOICE_LABEL = "Hungarian | Anna | medium"
PREVIEW_OUTPUT_PATH = Path.cwd() / "output" / "read-aloud-preview.mp3"
PAUSE_MS = 300
MAX_CHARS_PER_CHUNK = 280
READ_ALOUD_LINE_TAG = "read_aloud_line"
SUPPORTED_OUTPUT_FORMATS = {
    ".mp3": "MP3",
    ".ogg": "OGG",
    ".wav": "WAV",
}
MP3_QUALITY_PRESETS = {
    "64 kbps": {"bitrate": "64k"},
    "128 kbps": {"bitrate": "128k"},
    "192 kbps (recommended)": {"bitrate": "192k"},
    "256 kbps": {"bitrate": "256k"},
    "320 kbps": {"bitrate": "320k"},
}
OGG_QUALITY_PRESETS = {
    "Low (q1)": {"quality_params": ["-q:a", "1"]},
    "Medium (q3)": {"quality_params": ["-q:a", "3"]},
    "High (q5)": {"quality_params": ["-q:a", "5"]},
    "Very High (q8)": {"quality_params": ["-q:a", "8"]},
    "Maximum (q10)": {"quality_params": ["-q:a", "10"]},
}
MAX_MERGE_CHUNKS = 500
PREVIEW_FILE_GLOB = "read-aloud-preview-*.wav"
MIN_CHAPTER_WORDS = 30

CEFR_PRESETS: dict[str, dict[str, object]] = {
    "A1": {"level": 0, "count": 10, "max_length": 60, "plural_chance": 0.0},
    "A2": {"level": 1, "count": 15, "max_length": 80, "plural_chance": 0.3},
    "B1": {"level": 2, "count": 20, "max_length": 100, "plural_chance": 0.4},
    "B2": {"level": 3, "count": 25, "max_length": 120, "plural_chance": 0.5},
}

DEFAULT_LANG_LEARNING_SETTINGS: dict[str, object] = {
    "preset": "A2",
    "language": "pt",
    "level": 1,
    "count": 15,
    "max_length": 80,
    "plural_chance": 0.3,
    "seed": None,
    "top_n": None,
    "base_word": None,
    "base_word_count": 10,
    "base_template": None,
    "vary_role": None,
    "vary_words": None,
    "pair_pause_ms": 2000,
    "auto_speak": False,
    "show_translations": True,
}

HEADING_PATTERNS: list[tuple[str, int]] = [
    (r'(?m)^\s*(?:Chapter|CHAPTER)\s+(\d+|[IVXLCDM]+)\b', 1),
    (r'(?m)^\s*(\d+)\.\s+(?:[A-Z][\w\s]{3,})$', 1),
    (r'(?m)^\s*(?:Part|PART)\s+(\d+|[IVXLCDM]+)\b', 1),
    (r'(?m)^\s*(?:Section|SECTION)\s+(\d+)\b', 2),
    (r'(?m)^\s*[IVXLCDM]+\.\s', 1),
]
FONT_BODY = "Segoe UI" if sys.platform == "win32" else "DejaVu Sans"
FONT_MONO = "Consolas" if sys.platform == "win32" else "Liberation Mono"

def sanitize_filename(name: str) -> str:
    return re.sub(r'[<>:"/\\|?*]', '_', name).strip().rstrip('.')


SURFACE_BG = "#f4f6fb"
CARD_BG = "#ffffff"
ACCENT = "#2563eb"
ACCENT_ACTIVE = "#1d4ed8"
MUTED_TEXT = "#52607a"
TEXT_BG = "#fbfcfe"
TEXT_BORDER = "#d7deea"
READ_ALOUD_HIGHLIGHT = "#fff3bf"

THEMES: dict[str, dict[str, str]] = {
    "light": {
        "surface_bg": "#f4f6fb",
        "card_bg": "#ffffff",
        "accent": "#2563eb",
        "accent_active": "#1d4ed8",
        "muted_text": "#52607a",
        "text_bg": "#fbfcfe",
        "text_border": "#d7deea",
        "log_bg": "#eef2ff",
        "log_fg": "#172554",
        "read_aloud_highlight": "#fff3bf",
        "header_fg": "#101828",
        "label_fg": "#0f172a",
        "button_bg": "#ffffff",
        "button_fg": "#0f172a",
        "button_active_bg": "#eef2ff",
        "notebook_tab_bg": "#ffffff",
        "notebook_tab_fg": "#0f172a",
        "notebook_tab_selected_bg": "#2563eb",
        "notebook_tab_selected_fg": "#ffffff",
    },
    "dark": {
        "surface_bg": "#1e1e2e",
        "card_bg": "#282838",
        "accent": "#3b82f6",
        "accent_active": "#2563eb",
        "muted_text": "#9ca3af",
        "text_bg": "#1e1e2e",
        "text_border": "#374151",
        "log_bg": "#0f0f1a",
        "log_fg": "#e5e7eb",
        "read_aloud_highlight": "#3d2b00",
        "header_fg": "#e5e7eb",
        "label_fg": "#e5e7eb",
        "button_bg": "#282838",
        "button_fg": "#e5e7eb",
        "button_active_bg": "#374151",
        "notebook_tab_bg": "#282838",
        "notebook_tab_fg": "#e5e7eb",
        "notebook_tab_selected_bg": "#3b82f6",
        "notebook_tab_selected_fg": "#ffffff",
    },
    "high_contrast": {
        "surface_bg": "#ffffff",
        "card_bg": "#ffffff",
        "accent": "#0037a6",
        "accent_active": "#001f61",
        "muted_text": "#1f2937",
        "text_bg": "#ffffff",
        "text_border": "#111827",
        "log_bg": "#f3f4f6",
        "log_fg": "#111827",
        "read_aloud_highlight": "#fff176",
        "header_fg": "#000000",
        "label_fg": "#000000",
        "button_bg": "#ffffff",
        "button_fg": "#000000",
        "button_active_bg": "#dbeafe",
        "notebook_tab_bg": "#ffffff",
        "notebook_tab_fg": "#000000",
        "notebook_tab_selected_bg": "#0037a6",
        "notebook_tab_selected_fg": "#000000",
    },
}

CURRENT_THEME = "light"


def resolve_theme(theme_name: str) -> str:
    if theme_name != "system":
        return theme_name
    try:
        if sys.platform == "win32":
            import winreg
            key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, r"Software\Microsoft\Windows\CurrentVersion\Themes\Personalize")
            value, _ = winreg.QueryValueEx(key, "AppsUseLightTheme")
            winreg.CloseKey(key)
            return "light" if value == 1 else "dark"
        elif sys.platform == "darwin":
            result = subprocess.run(["defaults", "read", "-g", "AppleInterfaceStyle"], capture_output=True, text=True)
            return "dark" if "Dark" in result.stdout else "light"
        else:
            gtk_theme = os.environ.get("GTK_THEME", "").lower()
            if "dark" in gtk_theme:
                return "dark"
            return "light"
    except Exception:
        return "light"


def apply_theme(theme_key: str, root: Tk | None = None) -> None:
    global SURFACE_BG, CARD_BG, ACCENT, ACCENT_ACTIVE, MUTED_TEXT, TEXT_BG, TEXT_BORDER, READ_ALOUD_HIGHLIGHT
    global CURRENT_THEME
    colors = THEMES[theme_key]
    SURFACE_BG = colors["surface_bg"]
    CARD_BG = colors["card_bg"]
    ACCENT = colors["accent"]
    ACCENT_ACTIVE = colors["accent_active"]
    MUTED_TEXT = colors["muted_text"]
    TEXT_BG = colors["text_bg"]
    TEXT_BORDER = colors["text_border"]
    READ_ALOUD_HIGHLIGHT = colors["read_aloud_highlight"]
    CURRENT_THEME = theme_key

    style = ttk.Style()
    style.configure("TFrame", background=SURFACE_BG)
    style.configure("HeaderPanel.TFrame", background=CARD_BG)
    style.configure("Toolbar.TFrame", background=SURFACE_BG)
    style.configure("Sidebar.TFrame", background=CARD_BG)
    style.configure("TLabel", background=SURFACE_BG, foreground=colors["label_fg"])
    style.configure("Header.TLabel", background=CARD_BG, foreground=colors["header_fg"], font=(FONT_BODY, 20, "bold"))
    style.configure("HeroIcon.TLabel", background=CARD_BG, foreground=ACCENT, font=(FONT_BODY, 22))
    style.configure("Subtle.TLabel", background=CARD_BG, foreground=MUTED_TEXT, font=(FONT_BODY, 10))
    style.configure("Hint.TLabel", background=SURFACE_BG, foreground=MUTED_TEXT, font=(FONT_BODY, 9))
    style.configure(
        "TLabelframe",
        background=CARD_BG,
        bordercolor=TEXT_BORDER,
        relief="solid",
        borderwidth=1,
    )
    style.configure("TLabelframe.Label", background=CARD_BG, foreground=colors["label_fg"], font=(FONT_BODY, 10, "bold"))
    style.configure(
        "TButton",
        padding=(12, 8),
        font=(FONT_BODY, 9, "bold"),
        background=colors["button_bg"],
        foreground=colors["button_fg"],
        bordercolor=TEXT_BORDER,
        focusthickness=0,
    )
    style.map("TButton", background=[("active", colors["button_active_bg"])], bordercolor=[("active", ACCENT)])
    style.configure(
        "Accent.TButton",
        background=ACCENT,
        foreground="#ffffff" if theme_key != "high_contrast" else "#000000",
        bordercolor=ACCENT,
    )
    style.map(
        "Accent.TButton",
        background=[("active", ACCENT_ACTIVE)],
        bordercolor=[("active", ACCENT_ACTIVE)],
        foreground=[("active", "#ffffff" if theme_key != "high_contrast" else "#000000")],
    )
    style.configure("TEntry", fieldbackground=CARD_BG, bordercolor=TEXT_BORDER, padding=6)
    style.configure("TCombobox", fieldbackground=CARD_BG, bordercolor=TEXT_BORDER, padding=6)
    style.configure("TMenubutton", padding=(10, 7), background=colors["button_bg"], foreground=colors["button_fg"])
    style.configure("TCheckbutton", background=SURFACE_BG, foreground=colors["label_fg"])
    style.configure("Treeview", background=CARD_BG, fieldbackground=CARD_BG, foreground=colors["label_fg"])
    style.configure("Treeview.Heading", background=colors["button_active_bg"], foreground=colors["label_fg"])
    style.configure("TSeparator", background=TEXT_BORDER)
    style.configure("TNotebook", background=SURFACE_BG, bordercolor=TEXT_BORDER)
    style.configure("TNotebook.Tab", background=colors["notebook_tab_bg"], foreground=colors["notebook_tab_fg"], bordercolor=TEXT_BORDER, padding=(12, 8))
    style.map("TNotebook.Tab", background=[("selected", colors["notebook_tab_selected_bg"])], foreground=[("selected", colors["notebook_tab_selected_fg"])])

    if root is not None:
        root.configure(bg=SURFACE_BG)
        root.update_idletasks()


def sentence_split(text: str) -> list[str]:
    parts = re.split(r"(?<=[.!?;:])\s+|\n{2,}", text.strip(), flags=re.MULTILINE)
    return [part.strip() for part in parts if part.strip()]


def split_long_piece(piece: str, max_chars: int) -> list[str]:
    if len(piece) <= max_chars:
        return [piece]

    chunks: list[str] = []
    clauses = re.split(r"(?<=[,])\s+|(?<=\))\s+", piece)
    current = ""
    for clause in clauses:
        clause = clause.strip()
        if not clause:
            continue
        candidate = f"{current} {clause}".strip()
        if current and len(candidate) > max_chars:
            chunks.append(current)
            current = clause
        elif len(clause) > max_chars:
            words = clause.split()
            word_chunk = ""
            for word in words:
                word_candidate = f"{word_chunk} {word}".strip()
                if word_chunk and len(word_candidate) > max_chars:
                    chunks.append(word_chunk)
                    word_chunk = word
                else:
                    word_chunk = word_candidate
            if word_chunk:
                if current:
                    chunks.append(current)
                    current = ""
                chunks.append(word_chunk)
        else:
            current = candidate
    if current:
        chunks.append(current)
    return chunks


def chunk_text(text: str, max_chars: int = MAX_CHARS_PER_CHUNK) -> list[str]:
    sentences = sentence_split(text)
    if not sentences:
        return []

    chunks: list[str] = []
    current = ""
    for sentence in sentences:
        for piece in split_long_piece(sentence, max_chars):
            candidate = f"{current} {piece}".strip()
            if current and len(candidate) > max_chars:
                chunks.append(current)
                current = piece
            else:
                current = candidate
    if current:
        chunks.append(current)
    return chunks


@dataclass
class TextChunk:
    text: str
    start: int
    end: int


def find_word_start_offset(text: str, offset: int) -> int | None:
    clamped = max(0, min(offset, len(text)))
    for match in re.finditer(r"\S+", text):
        if match.start() <= clamped < match.end():
            return match.start()
        if match.start() > clamped:
            return match.start()
    return None


def chunk_text_with_offsets(
    text: str,
    max_chars: int = MAX_CHARS_PER_CHUNK,
    start_offset: int = 0,
) -> list[TextChunk]:
    matches = list(re.finditer(r"\S+", text))
    if not matches:
        return []

    effective_start = find_word_start_offset(text, start_offset)
    if effective_start is None:
        return []

    start_index = 0
    for index, match in enumerate(matches):
        if match.start() >= effective_start:
            start_index = index
            break

    chunks: list[TextChunk] = []
    chunk_start = matches[start_index].start()
    chunk_end = matches[start_index].end()

    for match in matches[start_index + 1:]:
        candidate = text[chunk_start:match.end()].strip()
        if candidate and len(candidate) > max_chars:
            chunk_text_value = text[chunk_start:chunk_end].strip()
            if chunk_text_value:
                chunks.append(TextChunk(text=chunk_text_value, start=chunk_start, end=chunk_end))
            chunk_start = match.start()
        chunk_end = match.end()

    chunk_text_value = text[chunk_start:chunk_end].strip()
    if chunk_text_value:
        chunks.append(TextChunk(text=chunk_text_value, start=chunk_start, end=chunk_end))
    return chunks


def output_format_for_path(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_OUTPUT_FORMATS:
        allowed = ", ".join(SUPPORTED_OUTPUT_FORMATS)
        raise ValueError(f"Unsupported output format '{suffix or '(none)'}'. Choose one of: {allowed}.")
    return suffix[1:]


def export_audio_segment(
    audio: AudioSegment,
    output_file: Path,
    bitrate: str | None = None,
    quality_params: list[str] | None = None,
) -> None:
    output_file.parent.mkdir(parents=True, exist_ok=True)
    output_format = output_format_for_path(output_file)
    export_kwargs: dict[str, object] = {"format": output_format}

    if quality_params is not None:
        export_kwargs["parameters"] = ["-ar", "44100"] + quality_params
    elif bitrate is not None:
        export_kwargs["bitrate"] = bitrate
        export_kwargs["parameters"] = ["-ar", "44100"]
    elif output_format in {"mp3", "ogg"}:
        export_kwargs["bitrate"] = "192k"
        export_kwargs["parameters"] = ["-ar", "44100"]
    else:
        export_kwargs["parameters"] = ["-ar", "44100"]

    audio.export(output_file, **export_kwargs)


class DocumentExtractor:
    SUPPORTED = {
        ".docx": "DOCX",
        ".odt": "ODT",
        ".pdf": "PDF",
        ".epub": "EPUB",
        ".mobi": "MOBI",
    }

    @staticmethod
    def extract_text(filepath: Path, from_page: int | None = None, to_page: int | None = None) -> str:
        suffix = filepath.suffix.lower()
        if suffix == ".docx":
            return DocumentExtractor._extract_docx(filepath)
        if suffix == ".odt":
            return DocumentExtractor._extract_odt(filepath)
        if suffix == ".pdf":
            return DocumentExtractor._extract_pdf(filepath, from_page=from_page, to_page=to_page)
        if suffix == ".epub":
            return DocumentExtractor._extract_epub(filepath)
        if suffix == ".mobi":
            return DocumentExtractor._extract_mobi(filepath)
        raise ValueError(f"Unsupported document format: {suffix}")

    @staticmethod
    def _extract_docx(filepath: Path) -> str:
        try:
            import docx
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
        doc = docx.Document(str(filepath))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    @staticmethod
    def _extract_odt(filepath: Path) -> str:
        try:
            from odf import teletype
            from odf import text as odf_text
            from odf.opendocument import load
        except ImportError:
            raise RuntimeError("odfpy not installed. Run: pip install odfpy")
        doc = load(str(filepath))
        paragraphs = []
        for elem in doc.getElementsByType(odf_text.P):
            content = teletype.extractText(elem).strip()
            if content:
                paragraphs.append(content)
        return "\n\n".join(paragraphs)

    @staticmethod
    def _extract_pdf(filepath: Path, from_page: int | None = None, to_page: int | None = None) -> str:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise RuntimeError("pypdf not installed. Run: pip install pypdf")
        reader = PdfReader(str(filepath))
        total_pages = len(reader.pages)
        start = (from_page or 1) - 1
        end = to_page if to_page else total_pages
        start = max(0, min(start, total_pages - 1))
        end = max(start + 1, min(end, total_pages))
        pages = []
        for page in reader.pages[start:end]:
            text = page.extract_text()
            if text:
                text = re.sub(r"\n{3,}", "\n\n", text)
                text = re.sub(r" {2,}", " ", text)
                pages.append(text.strip())
        return "\n\n".join(pages)

    @staticmethod
    def _extract_epub(filepath: Path) -> str:
        try:
            import ebooklib
            from bs4 import BeautifulSoup
            from ebooklib import epub
        except ImportError:
            raise RuntimeError("ebooklib and/or beautifulsoup4 not installed. Run: pip install ebooklib beautifulsoup4")
        book = epub.read_epub(str(filepath))
        chapters = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), "html.parser")
            text = soup.get_text(separator="\n")
            text = re.sub(r"\n{3,}", "\n\n", text)
            if text.strip():
                chapters.append(text.strip())
        return "\n\n".join(chapters)

    @staticmethod
    def _extract_mobi(filepath: Path) -> str:
        try:
            from mobi import extract
        except ImportError:
            raise RuntimeError("mobi not installed. Run: pip install mobi")
        import shutil as _shutil_mod

        temp_dir = None
        try:
            temp_dir, extracted_path = extract(str(filepath))
            temp_dir = Path(temp_dir)
            extracted = Path(extracted_path)

            if extracted.suffix.lower() == ".epub":
                text = DocumentExtractor._extract_epub(extracted)
            elif extracted.suffix.lower() == ".html":
                try:
                    from bs4 import BeautifulSoup
                except ImportError:
                    raise RuntimeError(
                        "beautifulsoup4 not installed for MOBI HTML extraction. "
                        "Run: pip install beautifulsoup4"
                    )
                soup = BeautifulSoup(extracted.read_text(encoding="utf-8", errors="replace"), "html.parser")
                text = soup.get_text(separator="\n")
                text = re.sub(r"\n{3,}", "\n\n", text)
            else:
                text = extracted.read_text(encoding="utf-8", errors="replace")

            if not text.strip():
                raise RuntimeError("No text extracted from MOBI file.")
            return text
        except RuntimeError:
            raise
        except Exception as exc:
            raise RuntimeError(f"MOBI extraction failed: {exc}") from exc
        finally:
            if temp_dir is not None:
                try:
                    _shutil_mod.rmtree(str(temp_dir), ignore_errors=True)
                except Exception:
                    pass

    @staticmethod
    def extract_chapters(
        filepath: Path,
        min_level: str = "all",
        from_page: int | None = None,
        to_page: int | None = None,
    ) -> list[tuple[str, str]]:
        suffix = filepath.suffix.lower()
        if suffix == ".docx":
            return DocumentExtractor._extract_docx_chapters(filepath, min_level)
        if suffix == ".odt":
            return DocumentExtractor._extract_odt_chapters(filepath, min_level)
        if suffix == ".pdf":
            return DocumentExtractor._extract_pdf_chapters(filepath, min_level, from_page=from_page, to_page=to_page)
        if suffix == ".epub":
            return DocumentExtractor._extract_epub_chapters(filepath, min_level)
        if suffix == ".mobi":
            return DocumentExtractor._extract_mobi_chapters(filepath, min_level)
        return [(("", DocumentExtractor.extract_text(filepath, from_page=from_page, to_page=to_page)))]

    @staticmethod
    def _heading_level_for_setting(min_level: str) -> int | None:
        if min_level == "h1":
            return 1
        if min_level == "h1-h2":
            return 2
        if min_level == "h1-h3":
            return 3
        return None

    @staticmethod
    def _heading_level_from_style(style_name: str) -> int | None:
        if not style_name:
            return None
        match = re.match(r'[Hh]eading\s*(\d+)', style_name)
        if match:
            return int(match.group(1))
        return None

    @staticmethod
    def _extract_docx_chapters(filepath: Path, min_level: str = "all") -> list[tuple[str, str]]:
        try:
            import docx
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
        doc = docx.Document(str(filepath))
        max_level = DocumentExtractor._heading_level_for_setting(min_level)

        chapters: list[tuple[str, str]] = []
        current_title = ""
        current_content: list[str] = []

        for p in doc.paragraphs:
            heading_level = DocumentExtractor._heading_level_from_style(p.style.name if p.style else "")
            text = p.text.strip()
            if not text:
                continue

            if heading_level is not None and (max_level is None or heading_level <= max_level):
                if current_content:
                    chapters.append((current_title, "\n\n".join(current_content)))
                current_title = text
                current_content = []
            else:
                current_content.append(text)

        if current_content:
            chapters.append((current_title, "\n\n".join(current_content)))

        if not chapters:
            return [("", DocumentExtractor._extract_docx(filepath))]

        return chapters

    @staticmethod
    def _extract_odt_chapters(filepath: Path, min_level: str = "all") -> list[tuple[str, str]]:
        try:
            from odf import teletype
            from odf.opendocument import load
        except ImportError:
            raise RuntimeError("odfpy not installed. Run: pip install odfpy")
        doc = load(str(filepath))
        max_level = DocumentExtractor._heading_level_for_setting(min_level)

        chapters: list[tuple[str, str]] = []
        current_title = ""
        current_content: list[str] = []

        body = getattr(doc, "body", None)
        if body is None:
            return [("", DocumentExtractor._extract_odt(filepath))]

        children = getattr(body, "childNodes", [])
        if not children:
            return [("", DocumentExtractor._extract_odt(filepath))]

        text_children: list = []
        for body_child in children:
            tag = getattr(body_child, "tagName", "")
            if tag in ("text:h", "text:p"):
                text_children = children
                break
            nested = getattr(body_child, "childNodes", [])
            if nested:
                text_children = nested
                break

        if not text_children:
            text_children = children

        has_headings = False

        for child in text_children:
            tag_name = getattr(child, "tagName", "")
            text_content = teletype.extractText(child).strip()

            if tag_name == "text:h":
                has_headings = True
                outline_level = int(child.getAttribute("outlinelevel") or "1")
                if max_level is None or outline_level <= max_level:
                    if current_content:
                        chapters.append((current_title, "\n\n".join(current_content)))
                    current_title = text_content
                    current_content = []
                    continue
                current_content.append(text_content)
            elif tag_name == "text:p":
                if text_content:
                    current_content.append(text_content)

        if current_content:
            chapters.append((current_title, "\n\n".join(current_content)))

        if not has_headings:
            return [("", DocumentExtractor._extract_odt(filepath))]

        return chapters

    @staticmethod
    def _extract_pdf_chapters(
        filepath: Path,
        min_level: str = "all",
        from_page: int | None = None,
        to_page: int | None = None,
    ) -> list[tuple[str, str]]:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise RuntimeError("pypdf not installed. Run: pip install pypdf")
        reader = PdfReader(str(filepath))

        if reader.outline:
            chapters = DocumentExtractor._extract_pdf_chapters_from_outline(
                reader, from_page=from_page, to_page=to_page
            )
            if len(chapters) >= 2:
                return chapters

        return DocumentExtractor._extract_pdf_chapters_by_pattern(filepath, from_page=from_page, to_page=to_page)

    @staticmethod
    def _extract_pdf_chapters_from_outline(
        reader,
        from_page: int | None = None,
        to_page: int | None = None,
    ) -> list[tuple[str, str]]:
        total_pages = len(reader.pages)
        user_start = max(0, (from_page or 1) - 1)
        user_end = min(total_pages, to_page) if to_page else total_pages
        user_start = max(0, min(user_start, total_pages - 1))
        user_end = max(user_start + 1, min(user_end, total_pages))

        pages_text: dict[int, str] = {}
        for page_num in range(user_start, user_end):
            page = reader.pages[page_num]
            text = page.extract_text()
            if text:
                text = re.sub(r"\n{3,}", "\n\n", text)
                text = re.sub(r" {2,}", " ", text)
                pages_text[page_num] = text.strip()

        outline_items: list[tuple[str, int]] = []

        def walk_outline(items, depth=0):
            for item in items:
                if isinstance(item, list):
                    walk_outline(item, depth + 1)
                else:
                    title = getattr(item, "title", "")
                    if not title:
                        continue
                    title = title.strip()
                    try:
                        page_num = reader.get_destination_page_number(item)
                    except Exception:
                        continue
                    outline_items.append((title, page_num))

        walk_outline(reader.outline)

        if not outline_items:
            return []

        chapters: list[tuple[str, str]] = []
        for i, (title, start_page) in enumerate(outline_items):
            end_page = outline_items[i + 1][1] if i + 1 < len(outline_items) else len(reader.pages)
            chapter_pages = []
            for pg in range(start_page, min(end_page, len(reader.pages))):
                if pg in pages_text:
                    chapter_pages.append(pages_text[pg])
            content = "\n\n".join(chapter_pages).strip()
            if content:
                chapters.append((title, content))

        return chapters

    @staticmethod
    def _extract_pdf_chapters_by_pattern(
        filepath: Path,
        from_page: int | None = None,
        to_page: int | None = None,
    ) -> list[tuple[str, str]]:
        full_text = DocumentExtractor._extract_pdf(filepath, from_page=from_page, to_page=to_page)
        if not full_text.strip():
            return [("", "")]

        split_points: list[tuple[int, str]] = []
        for pattern, _level in HEADING_PATTERNS:
            for match in re.finditer(pattern, full_text):
                pos = match.start()
                if all(abs(pos - existing) > 5 for existing, _ in split_points):
                    split_points.append((pos, match.group().strip()))
            if split_points:
                break

        if not split_points:
            return [("", full_text)]

        split_points.sort(key=lambda x: x[0])

        chapters: list[tuple[str, str]] = []
        for i, (pos, title) in enumerate(split_points):
            end = split_points[i + 1][0] if i + 1 < len(split_points) else len(full_text)
            content = full_text[pos:end].strip()
            if content:
                chapters.append((title, content))

        if not chapters:
            return [("", full_text)]

        return chapters

    @staticmethod
    def _extract_epub_chapters(filepath: Path, min_level: str = "all") -> list[tuple[str, str]]:
        try:
            import ebooklib
            from bs4 import BeautifulSoup
            from ebooklib import epub
        except ImportError:
            raise RuntimeError("ebooklib and/or beautifulsoup4 not installed. Run: pip install ebooklib beautifulsoup4")
        book = epub.read_epub(str(filepath))
        max_level = DocumentExtractor._heading_level_for_setting(min_level)

        chapters: list[tuple[str, str]] = []
        current_title = ""
        current_content: list[str] = []
        has_headings = False

        def _detect_level(tag_name: str) -> int | None:
            if tag_name and tag_name.startswith("h") and len(tag_name) == 2:
                try:
                    return int(tag_name[1])
                except ValueError:
                    return None
            return None

        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), "html.parser")

            for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p"], recursive=False):
                tag_name = getattr(element, "name", "")
                text_content = element.get_text(separator="\n").strip()
                if not text_content:
                    continue

                heading_level = _detect_level(tag_name)
                if heading_level is not None and (max_level is None or heading_level <= max_level):
                    has_headings = True
                    if current_content:
                        chapters.append((current_title, "\n\n".join(current_content)))
                    current_title = text_content
                    current_content = []
                else:
                    current_content.append(text_content)

            text = soup.get_text(separator="\n")
            text = re.sub(r"\n{3,}", "\n\n", text)
            if text.strip():
                for elem in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"], recursive=True):
                    if _detect_level(getattr(elem, "name", "")):
                        has_headings = True
                        break

        if current_content:
            chapters.append((current_title, "\n\n".join(current_content)))

        if not has_headings or not chapters:
            return [("", DocumentExtractor._extract_epub(filepath))]

        return chapters

    @staticmethod
    def _extract_mobi_chapters(filepath: Path, min_level: str = "all") -> list[tuple[str, str]]:
        try:
            from mobi import extract
        except ImportError:
            raise RuntimeError("mobi not installed. Run: pip install mobi")
        import shutil as _shutil_mod

        temp_dir = None
        try:
            temp_dir, extracted_path = extract(str(filepath))
            temp_dir = Path(temp_dir)
            extracted = Path(extracted_path)

            if extracted.suffix.lower() == ".epub":
                result = DocumentExtractor._extract_epub_chapters(extracted, min_level)
            else:
                result = [("", DocumentExtractor._extract_mobi(filepath))]

            return result
        except Exception as exc:
            raise RuntimeError(f"MOBI extraction failed: {exc}") from exc
        finally:
            if temp_dir is not None:
                try:
                    _shutil_mod.rmtree(str(temp_dir), ignore_errors=True)
                except Exception:
                    pass


@dataclass
class SynthesisRequest:
    text: str
    language: str
    output_file: Path
    engine: str
    piper_voice_label: str
    piper_voice_code: str
    speaker_name: str
    speaker_wav: str
    speed: float = 1.0


class SessionState(str, Enum):
    IDLE = "idle"
    GENERATING = "generating"
    SYNTHESIZING = "synthesizing"
    PLAYING = "playing"
    PAUSED = "paused"
    STOPPING = "stopping"
    FAILED = "failed"


@dataclass(frozen=True)
class LanguageLearningAvailability:
    available: bool
    message: str
    wordlist_dir: Path | None = None


@dataclass(frozen=True)
class LanguageSessionConfig:
    language: str
    engine: str
    speed: float
    pair_pause_ms: int
    show_translations: bool


def language_learning_availability() -> LanguageLearningAvailability:
    """Validate the optional companion before a wizard tries to use it."""
    try:
        import importlib

        package = importlib.import_module("language_practice")
        importlib.import_module("language_practice.languages.pt")
        importlib.import_module("language_practice.languages.es")
    except Exception:
        return LanguageLearningAvailability(
            False,
            "Language Learning needs the optional language-practice companion. "
            "Install it with: pip install -e ../language-practice",
        )

    candidates: list[Path] = []
    configured = os.environ.get("LANGUAGE_PRACTICE_DATA_DIR")
    if configured:
        candidates.append(Path(configured))
    package_file = getattr(package, "__file__", None)
    if package_file:
        # Editable installs keep `src/language_practice`; its project root owns data/.
        candidates.append(Path(package_file).resolve().parents[2] / "data")
    candidates.append(Path.cwd().parent / "language-practice" / "data")

    for data_dir in candidates:
        if (data_dir / "pt" / "wordlist.md").is_file() and (data_dir / "es" / "wordlist.md").is_file():
            return LanguageLearningAvailability(True, "Language Learning is ready.", data_dir)
    return LanguageLearningAvailability(
        False,
        "The language-practice companion is installed but its PT/ES wordlists "
        "were not found. Set LANGUAGE_PRACTICE_DATA_DIR or use an editable install.",
    )


@dataclass(frozen=True)
class ChapterEntry:
    source_path: Path
    index: int
    title: str
    content: str
    word_count: int


def get_piper_voice_metadata(label: str) -> dict[str, str]:
    metadata = PIPER_VOICE_OPTIONS.get(label)
    if metadata is None:
        return PIPER_VOICE_OPTIONS[DEFAULT_PIPER_VOICE_LABEL]
    return metadata


def label_for_piper_voice(voice_code: str, voice_info: dict | None = None) -> str:
    if voice_info is not None:
        language = voice_info.get("language", {})
        language_name = language.get("name_english", voice_code.split("-")[0])
        country_name = language.get("country_english", "")
        region = f" {country_name}" if country_name else ""
        voice_name = str(voice_info.get("name", "")).replace("_", " ").title()
        quality = voice_info.get("quality", "")
        return f"{language_name}{region} | {voice_name} | {quality}"

    parts = voice_code.split("-")
    lang_code = parts[0]
    name = parts[1] if len(parts) > 1 else "voice"
    quality = parts[2] if len(parts) > 2 else "unknown"
    language_map = {
        "hu_HU": "Hungarian",
        "en_US": "English United States",
        "en_GB": "English Great Britain",
    }
    language_name = language_map.get(lang_code, lang_code)
    return f"{language_name} | {name.replace('_', ' ').title()} | {quality}"


DEFAULT_UI_SETTINGS: dict[str, object] = {
    "theme": "light",
    "sidebar_collapsed": False,
    "toolbar_compact": False,
    "font_size": 11,
}

DEFAULT_AUDIO_SETTINGS: dict[str, object] = {
    "default_format": "MP3",
    "mp3_quality": "192 kbps (recommended)",
    "ogg_quality": "High (q5)",
    "sample_rate": 44100,
}

DEFAULT_GENERAL_SETTINGS: dict[str, object] = {
    "default_engine": ENGINE_AUTO,
    "default_language": "hu",
    "auto_save_output_folder": True,
    "confirm_on_exit": True,
    "xtts_license_accepted": False,
}

def get_default_music_folder() -> Path:
    if sys.platform == "win32":
        import ctypes
        from ctypes import wintypes

        csidl = 13  # CSIDL_MYMUSIC
        buf = ctypes.create_unicode_buffer(wintypes.MAX_PATH)
        if ctypes.windll.shell32.SHGetFolderPathW(None, csidl, None, 0, buf) == 0:
            return Path(buf.value)
        return Path.home() / "Music"

    if sys.platform == "darwin":
        return Path.home() / "Music"

    user_dirs = Path.home() / ".config" / "user-dirs.dirs"
    if user_dirs.exists():
        for line in user_dirs.read_text().splitlines():
            if line.startswith("XDG_MUSIC_DIR="):
                dir_path = line.split("=", 1)[1].strip('"').strip("'")
                dir_path = dir_path.replace("$HOME", str(Path.home()))
                resolved = Path(dir_path).expanduser()
                if resolved.exists():
                    return resolved
    return Path.home() / "Music"


DEFAULT_PATHS_SETTINGS: dict[str, object] = {
    "piper_voice_dir": "voices/piper",
    "pocket_voice_dir": "voices/pocket",
    "output_folder": str(get_default_music_folder()),
}


def load_app_settings() -> dict:
    if not APP_SETTINGS_PATH.exists():
        return _default_settings_dict()
    try:
        settings = json.loads(APP_SETTINGS_PATH.read_text(encoding="utf-8"))
    except Exception as exc:
        import logging
        logging.warning(f"Corrupted settings.json, using defaults: {exc}")
        # Backup corrupted file
        try:
            backup_path = APP_SETTINGS_PATH.with_suffix(".json.bak")
            APP_SETTINGS_PATH.rename(backup_path)
        except Exception:
            pass
        settings = {}
    return normalize_app_settings(settings)


def _default_settings_dict() -> dict:
    """Return settings dict with all default sections populated."""
    return {
        "language_learning": DEFAULT_LANG_LEARNING_SETTINGS.copy(),
        "ui": DEFAULT_UI_SETTINGS.copy(),
        "audio": DEFAULT_AUDIO_SETTINGS.copy(),
        "general": DEFAULT_GENERAL_SETTINGS.copy(),
        "paths": DEFAULT_PATHS_SETTINGS.copy(),
    }


def normalize_learning_language(value: object) -> str:
    """Return a supported ISO language code for Language Learning settings."""
    if not isinstance(value, str):
        return "pt"
    normalized = language_code_from_display(value.strip())
    return normalized if normalized in {"pt", "es"} else "pt"


def normalize_app_settings(raw_settings: object) -> dict:
    """Merge untrusted persisted JSON with defaults and migrate old values."""
    raw = raw_settings if isinstance(raw_settings, dict) else {}
    settings = dict(raw)
    sections = {
        "language_learning": DEFAULT_LANG_LEARNING_SETTINGS,
        "ui": DEFAULT_UI_SETTINGS,
        "audio": DEFAULT_AUDIO_SETTINGS,
        "general": DEFAULT_GENERAL_SETTINGS,
        "paths": DEFAULT_PATHS_SETTINGS,
    }
    for name, defaults in sections.items():
        current = raw.get(name)
        section = dict(current) if isinstance(current, dict) else {}
        settings[name] = {**defaults, **section}

    # `system` was the old default and can open a dark window on first launch.
    # Preserve deliberate light/dark/high-contrast choices, but make legacy
    # system settings deterministic and light.
    theme = settings["ui"].get("theme")
    settings["ui"]["theme"] = theme if theme in THEMES and theme != "system" else "light"
    settings["language_learning"]["language"] = normalize_learning_language(
        settings["language_learning"].get("language")
    )
    settings["general"]["default_language"] = language_code_from_display(
        str(settings["general"].get("default_language", "hu"))
    )
    return settings


def save_app_settings(settings: dict) -> None:
    """Persist settings atomically while retaining the previous valid file."""
    if APP_SETTINGS_PATH.exists():
        try:
            APP_SETTINGS_PATH.with_suffix(".json.bak").write_text(
                APP_SETTINGS_PATH.read_text(encoding="utf-8"), encoding="utf-8"
            )
        except OSError:
            pass
    temp_path = APP_SETTINGS_PATH.with_suffix(".json.tmp")
    temp_path.write_text(json.dumps(settings, indent=2), encoding="utf-8")
    temp_path.replace(APP_SETTINGS_PATH)


def add_context_menu(widget) -> None:
    """Attach right-click context menu with Cut/Copy/Paste/Select All to a Text/Entry widget."""
    menu = Menu(widget, tearoff=0)
    menu.add_command(label="Cut", command=lambda: widget.event_generate("<<Cut>>"))
    menu.add_command(label="Copy", command=lambda: widget.event_generate("<<Copy>>"))
    menu.add_command(label="Paste", command=lambda: widget.event_generate("<<Paste>>"))
    menu.add_separator()
    menu.add_command(label="Select All", command=lambda: widget.event_generate("<<SelectAll>>"))

    def show_menu(event):
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()

    widget.bind("<Button-3>", show_menu)
    widget.bind("<Control-a>", lambda e: widget.event_generate("<<SelectAll>>"))


def make_validated_entry(parent, textvariable, validator="int", width=8, **kwargs):
    """Create an Entry with input validation for int/float."""
    if validator == "int":
        vcmd = (parent.register(lambda s: s == "" or s.lstrip("-").isdigit()), "%P")
    elif validator == "float":
        vcmd = (parent.register(lambda s: s == "" or (s.replace(".", "", 1).lstrip("-").isdigit() and s.count(".") <= 1)), "%P")
    else:
        vcmd = None
    entry = ttk.Entry(parent, textvariable=textvariable, width=width, validate="key", validatecommand=vcmd, **kwargs)
    add_context_menu(entry)
    return entry


def discover_local_piper_voices() -> dict[str, dict[str, str]]:
    options = dict(PIPER_VOICE_OPTIONS)
    if not PIPER_VOICE_DIR.exists():
        return options

    for model_path in sorted(PIPER_VOICE_DIR.glob("*.onnx")):
        config_path = model_path.with_suffix(".onnx.json")
        if not config_path.exists():
            continue

        voice_code = model_path.stem
        if any(metadata["code"] == voice_code for metadata in options.values()):
            continue
        label = label_for_piper_voice(voice_code)
        language_prefix = voice_code.split("-")[0]
        xtts_language = "hu" if language_prefix == "hu_HU" else "en"
        options.setdefault(
            label,
            {
                "code": voice_code,
                "xtts_language": xtts_language,
            },
        )

    return dict(sorted(options.items()))


def piper_model_path(voice_code: str) -> Path:
    return PIPER_VOICE_DIR / f"{voice_code}.onnx"


def piper_language_of_code(voice_code: str) -> str:
    """Extract the ISO language ('hu', 'en', 'fr'...) from a Piper voice code.

    Piper codes look like 'hu_HU-anna-medium' or 'fr_FR-siwis-low'; the language
    is the segment before the first underscore.
    """
    prefix = voice_code.split("-", 1)[0]
    return prefix.split("_", 1)[0].lower()


def piper_languages(options: dict[str, dict[str, str]]) -> set[str]:
    """Languages actually available in Piper, derived from installed voices."""
    return {piper_language_of_code(meta["code"]) for meta in options.values()}


def piper_voices_for_language(
    options: dict[str, dict[str, str]], language: str
) -> list[str]:
    """Piper voice labels whose voice speaks the given language."""
    return [
        label
        for label, meta in options.items()
        if piper_language_of_code(meta["code"]) == language
    ]


def engines_supporting_language(
    language: str, piper_options: dict[str, dict[str, str]]
) -> list[str]:
    """Concrete engines that can speak `language`, in preference order."""
    engines: list[str] = []
    if language in piper_languages(piper_options):
        engines.append(ENGINE_PIPER)
    if language in XTTS_LANGUAGES:
        engines.append(ENGINE_XTTS)
    if language in POCKET_LANGUAGES:
        engines.append(ENGINE_POCKET)
    return engines


def available_languages(piper_options: dict[str, dict[str, str]]) -> list[str]:
    """Union of every language any engine can speak, in display order."""
    langs = set(XTTS_LANGUAGES) | set(POCKET_LANGUAGES) | piper_languages(piper_options)
    ordered = [code for code in LANGUAGE_NAMES if code in langs]
    ordered += sorted(code for code in langs if code not in LANGUAGE_NAMES)
    return ordered


def language_display_name(code: str) -> str:
    return LANGUAGE_NAMES.get(code, code.upper())


def language_code_from_display(display: str) -> str:
    for code, name in LANGUAGE_NAMES.items():
        if name == display:
            return code
    return display.lower()


def pocket_default_voice(language: str) -> str:
    return POCKET_DEFAULT_VOICE_FOR_LANG.get(language, POCKET_DEFAULT_VOICE)


def is_pocket_default_voice(name: str) -> bool:
    """True when `name` is one of the per-language defaults (i.e. not a voice the
    user deliberately picked), so it is safe to swap when the language changes."""
    return not name or name in set(POCKET_DEFAULT_VOICE_FOR_LANG.values()) | {POCKET_DEFAULT_VOICE}


def select_engine(
    engine: str,
    language: str,
    has_reference_wav: bool,
    piper_options: dict[str, dict[str, str]],
) -> str:
    """Resolve the effective engine, honouring an explicit choice or resolving
    ``Auto`` to an engine that can actually speak ``language``.

    Auto prefers XTTS when a reference clip is supplied (for cloning), then
    Piper, then any remaining engine that supports the language. This prevents
    Auto from silently routing to an engine with no voice for the language.
    """
    if engine and engine != ENGINE_AUTO:
        return engine
    supported = engines_supporting_language(language, piper_options)
    if has_reference_wav and ENGINE_XTTS in supported:
        return ENGINE_XTTS
    if ENGINE_PIPER in supported:
        return ENGINE_PIPER
    if supported:
        return supported[0]
    return ENGINE_XTTS if has_reference_wav else ENGINE_PIPER


class XTTSService:
    def __init__(self, log: Callable[[str], None]) -> None:
        self._log = log
        self._tts = None
        self._sample_rate = 24000

    def ensure_loaded(self) -> None:
        if self._tts is not None:
            return
        self._log("Loading XTTS v2 model. First load can take a while.")
        try:
            import torch
            from TTS.api import TTS
        except Exception as exc:  # pragma: no cover - import path varies by machine
            raise RuntimeError(
                "Coqui TTS is not installed in this environment. Run the setup script first (`.\\setup.ps1` on Windows or `./setup.sh` on Linux/macOS)."
            ) from exc

        self._patch_torch_load(torch)
        device = "cuda" if torch.cuda.is_available() else "cpu"
        self._log(f"Using device: {device}")
        self._tts = TTS(MODEL_NAME).to(device)
        sample_rate = getattr(self._tts.synthesizer, "output_sample_rate", None)
        if isinstance(sample_rate, int):
            self._sample_rate = sample_rate
        self._log("Model is ready.")

    @staticmethod
    def _patch_torch_load(torch_module) -> None:
        current = torch_module.load
        if getattr(current, "_coqui_xtts_patched", False):
            return

        def patched_load(*args, **kwargs):
            kwargs.setdefault("weights_only", False)
            return current(*args, **kwargs)

        patched_load._coqui_xtts_patched = True
        torch_module.load = patched_load

    def iter_segments(self, request: SynthesisRequest, start_offset: int = 0):
        self.ensure_loaded()
        AudioSegment.converter = imageio_ffmpeg.get_ffmpeg_exe()
        chunks = chunk_text_with_offsets(request.text, start_offset=start_offset)
        if not chunks:
            raise ValueError("Text is empty after cleanup.")

        self._log(f"Prepared {len(chunks)} chunk(s) for synthesis.")

        kwargs = {"language": request.language}
        if request.speed != 1.0:
            kwargs["speed"] = request.speed
        if request.speaker_wav:
            kwargs["speaker_wav"] = request.speaker_wav
            self._log("Voice source: reference WAV")
        else:
            kwargs["speaker"] = request.speaker_name or DEFAULT_SPEAKER
            self._log(f"Voice source: built-in speaker '{kwargs['speaker']}'")

        for index, chunk in enumerate(chunks, start=1):
            self._log(f"Synthesizing chunk {index}/{len(chunks)}")
            wav = self._tts.tts(text=chunk.text, split_sentences=True, **kwargs)
            array = np.asarray(wav, dtype=np.float32)
            pcm = np.int16(np.clip(array, -1.0, 1.0) * 32767).tobytes()
            segment = AudioSegment(
                data=pcm,
                sample_width=2,
                frame_rate=self._sample_rate,
                channels=1,
            )
            yield chunk, segment

    def synthesize(self, request: SynthesisRequest) -> Path:
        chunks = list(self.iter_segments(request))
        combined = AudioSegment.silent(duration=0)
        for index, (_chunk, segment) in enumerate(chunks, start=1):
            combined += segment
            if index < len(chunks):
                combined += AudioSegment.silent(duration=PAUSE_MS)

        self._log(f"Exporting audio to {request.output_file}")
        export_audio_segment(combined, request.output_file)
        self._log("Finished.")
        return request.output_file


class PiperService:
    def __init__(self, log: Callable[[str], None]) -> None:
        self._log = log
        self._voices: dict[str, object] = {}

    def ensure_loaded(self, voice_code: str):
        voice = self._voices.get(voice_code)
        if voice is not None:
            return voice

        model_path = piper_model_path(voice_code)
        if not model_path.exists():
            raise RuntimeError(
                f"Piper voice '{voice_code}' is missing. Run the setup script to download it (`.\\setup.ps1` on Windows or `./setup.sh` on Linux/macOS)."
            )

        self._log(f"Loading Piper voice '{voice_code}'.")
        try:
            from piper.voice import PiperVoice
        except Exception as exc:
            raise RuntimeError(
                "Piper is not installed in this environment. Run the setup script first (`.\\setup.ps1` on Windows or `./setup.sh` on Linux/macOS)."
            ) from exc

        voice = PiperVoice.load(model_path, download_dir=PIPER_VOICE_DIR)
        self._voices[voice_code] = voice
        self._log("Piper voice is ready.")
        return voice

    def iter_segments(self, request: SynthesisRequest, start_offset: int = 0):
        voice_code = request.piper_voice_code
        voice = self.ensure_loaded(voice_code)
        AudioSegment.converter = imageio_ffmpeg.get_ffmpeg_exe()
        chunks = chunk_text_with_offsets(request.text, start_offset=start_offset)
        if not chunks:
            raise ValueError("Text is empty after cleanup.")

        self._log(f"Prepared {len(chunks)} chunk(s) for synthesis.")
        self._log(f"Voice source: Piper '{voice_code}'")

        for index, chunk in enumerate(chunks, start=1):
            self._log(f"Synthesizing chunk {index}/{len(chunks)}")
            segment = AudioSegment.silent(duration=0)
            from piper.config import SynthesisConfig

            syn_config = None
            if request.speed != 1.0:
                syn_config = SynthesisConfig(length_scale=1.0 / request.speed)
            for audio_chunk in voice.synthesize(chunk.text, syn_config=syn_config):
                segment += AudioSegment(
                    data=audio_chunk.audio_int16_bytes,
                    sample_width=audio_chunk.sample_width,
                    frame_rate=audio_chunk.sample_rate,
                    channels=audio_chunk.sample_channels,
                )
            yield chunk, segment

    def synthesize(self, request: SynthesisRequest) -> Path:
        chunks = list(self.iter_segments(request))
        combined = AudioSegment.silent(duration=0)
        for index, (_chunk, segment) in enumerate(chunks, start=1):
            combined += segment
            if index < len(chunks):
                combined += AudioSegment.silent(duration=PAUSE_MS)

        self._log(f"Exporting audio to {request.output_file}")
        export_audio_segment(combined, request.output_file)
        self._log("Finished.")
        return request.output_file


class PocketTTSService:
    def __init__(self, log: Callable[[str], None]) -> None:
        self._log = log
        self._model = None
        self._loaded_language = None
        self._voice_states: dict = {}
        self._sample_rate = 24000

    def _resolve_language(self, lang_code: str) -> str:
        resolved = POCKET_LANG_MAP.get(lang_code)
        if resolved is None:
            self._log(
                f"Pocket TTS does not support '{lang_code}'; falling back to English. "
                "For this language, use Piper or XTTS instead."
            )
            return "english"
        return resolved

    @staticmethod
    def _looks_like_reference(voice_source: str) -> bool:
        """True when the source is a file to clone from or a remote URL, rather
        than a built-in voice name."""
        if Path(voice_source).is_file():
            return True
        lowered = voice_source.lower()
        return lowered.startswith(("http://", "https://", "hf://"))

    def _validate_voice_source(self, voice_source: str) -> None:
        if self._looks_like_reference(voice_source):
            return
        if voice_source not in POCKET_PREDEFINED_VOICES:
            preview = ", ".join(POCKET_PREDEFINED_VOICES[:6])
            raise RuntimeError(
                f"'{voice_source}' is not a built-in Pocket TTS voice. "
                f"Choose a built-in voice (e.g. {preview}) or pick a reference "
                "audio file to clone a voice from."
            )

    @staticmethod
    def _voice_cache_path(voice_source: str) -> Path | None:
        source_path = Path(voice_source)
        if source_path.is_file():
            return None
        safe = re.sub(r"[^a-zA-Z0-9_-]", "_", voice_source)
        return POCKET_VOICE_DIR / f"{safe}.safetensors"

    def ensure_loaded(self, lang_code: str = "en") -> None:
        lang = self._resolve_language(lang_code)
        if self._model is not None and self._loaded_language == lang:
            return
        if self._model is not None:
            self._log(
                f"Language changed ({self._loaded_language} → {lang}), reloading model..."
            )
            self._voice_states.clear()
        self._log(f"Loading Pocket TTS model (language: {lang})...")
        try:
            from pocket_tts import TTSModel
        except Exception as exc:
            raise RuntimeError(
                "pocket-tts is not installed in this environment. "
                "Run the setup script first (`.\\setup.ps1` on Windows or `./setup.sh` on Linux/macOS)."
            ) from exc
        try:
            self._model = TTSModel.load_model(language=lang)
        except Exception as exc:
            raise RuntimeError(
                f"Could not load the Pocket TTS model for {lang}. The model is "
                "downloaded on first use, so check your internet connection and "
                f"try again. (Details: {exc})"
            ) from exc
        self._sample_rate = self._model.sample_rate
        self._loaded_language = lang
        self._log("Pocket TTS model is ready.")

    def _get_voice_state(self, voice_source: str) -> dict:
        if voice_source in self._voice_states:
            return self._voice_states[voice_source]
        self._validate_voice_source(voice_source)
        cache_path = self._voice_cache_path(voice_source)
        voice_arg = str(cache_path) if cache_path is not None and cache_path.exists() else voice_source
        self._log(f"Loading voice from: {voice_arg}")
        try:
            state = self._model.get_state_for_audio_prompt(voice_arg)
        except Exception as exc:
            raise RuntimeError(
                f"Could not load the Pocket TTS voice '{voice_source}'. "
                "Built-in voice names must match exactly; reference clips must be "
                f"a readable audio file. (Details: {exc})"
            ) from exc
        if cache_path is not None and not cache_path.exists():
            POCKET_VOICE_DIR.mkdir(parents=True, exist_ok=True)
            try:
                from pocket_tts import export_model_state
                export_model_state(state, cache_path)
            except Exception as exc:
                self._log(f"Could not cache voice '{voice_source}' for reuse: {exc}")
        self._voice_states[voice_source] = state
        return state

    def iter_segments(self, request: SynthesisRequest, start_offset: int = 0):
        self.ensure_loaded(request.language)
        AudioSegment.converter = imageio_ffmpeg.get_ffmpeg_exe()
        chunks = chunk_text_with_offsets(request.text, start_offset=start_offset)
        if not chunks:
            raise ValueError("Text is empty after cleanup.")
        if request.speaker_wav:
            voice_source = request.speaker_wav
        else:
            voice_source = request.speaker_name or POCKET_DEFAULT_VOICE
        voice_state = self._get_voice_state(voice_source)
        self._log(f"Prepared {len(chunks)} chunk(s) for synthesis.")
        for index, chunk in enumerate(chunks, start=1):
            self._log(f"Synthesizing chunk {index}/{len(chunks)}")
            audio_tensor = self._model.generate_audio(
                voice_state, chunk.text, copy_state=True
            )
            array = audio_tensor.cpu().numpy()
            pcm = np.int16(np.clip(array, -1.0, 1.0) * 32767).tobytes()
            segment = AudioSegment(
                data=pcm,
                sample_width=2,
                frame_rate=self._sample_rate,
                channels=1,
            )
            yield chunk, segment

    def synthesize(self, request: SynthesisRequest) -> Path:
        chunks = list(self.iter_segments(request))
        combined = AudioSegment.silent(duration=0)
        for index, (_chunk, segment) in enumerate(chunks, start=1):
            combined += segment
            if index < len(chunks):
                combined += AudioSegment.silent(duration=PAUSE_MS)
        self._log(f"Exporting audio to {request.output_file}")
        export_audio_segment(combined, request.output_file)
        self._log("Finished.")
        return request.output_file


class SynthesisCoordinator:
    def __init__(self, log: Callable[[str], None]) -> None:
        self._xtts = XTTSService(log)
        self._piper = PiperService(log)
        self._pocket = PocketTTSService(log)

    @staticmethod
    def resolve_engine(request: SynthesisRequest) -> str:
        return select_engine(
            request.engine,
            request.language,
            bool(request.speaker_wav),
            discover_local_piper_voices(),
        )

    def synthesize(self, request: SynthesisRequest) -> Path:
        resolved_engine = self.resolve_engine(request)
        if resolved_engine == ENGINE_PIPER:
            return self._piper.synthesize(request)
        if resolved_engine == ENGINE_POCKET:
            return self._pocket.synthesize(request)
        return self._xtts.synthesize(request)

    def iter_segments(self, request: SynthesisRequest, start_offset: int = 0):
        resolved_engine = self.resolve_engine(request)
        if resolved_engine == ENGINE_PIPER:
            yield from self._piper.iter_segments(request, start_offset=start_offset)
            return
        if resolved_engine == ENGINE_POCKET:
            yield from self._pocket.iter_segments(request, start_offset=start_offset)
            return
        yield from self._xtts.iter_segments(request, start_offset=start_offset)


class AudioPlayer:
    def __init__(self, log: Callable[[str], None]) -> None:
        self._log = log
        self._ready = False
        self._paused = False
        self._active = False

    def ensure_ready(self) -> None:
        if self._ready:
            return
        try:
            import pygame
        except Exception as exc:
            raise RuntimeError(
                "pygame is not installed. Run the setup script first (`.\\setup.ps1` on Windows or `./setup.sh` on Linux/macOS)."
            ) from exc

        if not pygame.mixer.get_init():
            pygame.mixer.init()
        self._ready = True

    def play(self, path: Path) -> None:
        self.ensure_ready()
        import pygame

        self._paused = False
        self._active = True
        pygame.mixer.music.load(str(path))
        pygame.mixer.music.play()
        self._log(f"Playing audio: {path.name}")

    def play_blocking(self, path: Path, stop_event: threading.Event) -> None:
        self.play(path)
        import pygame

        try:
            while True:
                if stop_event.is_set():
                    pygame.mixer.music.stop()
                    return
                if self._paused:
                    time.sleep(0.05)
                    continue
                if not pygame.mixer.music.get_busy():
                    return
                time.sleep(0.05)
        finally:
            self._paused = False
            self._active = False

    def pause(self) -> None:
        if not self._active:
            return
        self.ensure_ready()
        import pygame

        pygame.mixer.music.pause()
        self._paused = True
        self._log("Playback paused.")

    def resume(self) -> None:
        if not self._active:
            return
        self.ensure_ready()
        import pygame

        pygame.mixer.music.unpause()
        self._paused = False
        self._log("Playback resumed.")

    def stop(self, quiet: bool = False) -> None:
        if not self._ready:
            return
        import pygame

        was_active = self._active
        pygame.mixer.music.stop()
        self._paused = False
        self._active = False
        if was_active and not quiet:
            self._log("Playback stopped.")

    def is_active(self) -> bool:
        return self._active

    def is_paused(self) -> bool:
        return self._paused


class PiperVoiceWizard:
    def __init__(self, app: "App") -> None:
        self.app = app
        self.app = app
        self.window = Toplevel(app.root)
        self.window.transient(app.root)
        self.window.title("Piper Voice Wizard")
        self.window.geometry("980x620")
        self.window.minsize(900, 560)

        self.search = StringVar()
        self.quality = StringVar(value="all")
        self.installed_only = BooleanVar(value=False)
        self.status = StringVar(value="Loading Piper voice catalog...")

        self.catalog: dict[str, dict] = {}
        self.filtered_codes: list[str] = []
        self.downloading = False

        self._build_ui()
        self.search.trace_add("write", self.apply_filters)
        self.quality.trace_add("write", self.apply_filters)
        self.installed_only.trace_add("write", self.apply_filters)
        threading.Thread(target=self.load_catalog, daemon=True).start()

    def _build_ui(self) -> None:  # pragma: no cover
        frame = ttk.Frame(self.window, padding=12)
        frame.pack(fill="both", expand=True)
        frame.columnconfigure(0, weight=1)
        frame.rowconfigure(1, weight=1)

        controls = ttk.Frame(frame)
        controls.grid(row=0, column=0, sticky="ew")
        controls.columnconfigure(1, weight=1)

        ttk.Label(controls, text="Search").grid(row=0, column=0, sticky="w")
        ttk.Entry(controls, textvariable=self.search).grid(row=0, column=1, sticky="ew", padx=(8, 12))
        ttk.Label(controls, text="Quality").grid(row=0, column=2, sticky="w")
        ttk.Combobox(
            controls,
            textvariable=self.quality,
            values=["all", "x_low", "low", "medium", "high"],
            state="readonly",
            width=10,
        ).grid(row=0, column=3, sticky="w", padx=(8, 12))
        ttk.Checkbutton(controls, text="Installed only", variable=self.installed_only).grid(row=0, column=4, sticky="w")

        columns = ("label", "code", "quality", "installed")
        self.tree = ttk.Treeview(frame, columns=columns, show="headings", height=18)
        self.tree.heading("label", text="Voice")
        self.tree.heading("code", text="Code")
        self.tree.heading("quality", text="Quality")
        self.tree.heading("installed", text="Installed")
        self.tree.column("label", width=360)
        self.tree.column("code", width=220)
        self.tree.column("quality", width=90, anchor="center")
        self.tree.column("installed", width=90, anchor="center")
        self.tree.grid(row=1, column=0, sticky="nsew", pady=(12, 0))

        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=self.tree.yview)
        scrollbar.grid(row=1, column=1, sticky="ns", pady=(12, 0))
        self.tree.configure(yscrollcommand=scrollbar.set)

        actions = ttk.Frame(frame)
        actions.grid(row=2, column=0, sticky="ew", pady=(12, 0))
        ttk.Button(actions, text="↻ Refresh Catalog", command=self.refresh_catalog).pack(side="left")
        ttk.Button(actions, text="↓ Download Selected", style="Accent.TButton", command=self.download_selected).pack(side="right")
        ttk.Button(actions, text="★ Set As Default", command=self.set_selected_default).pack(side="right", padx=(0, 8))

        ttk.Label(frame, textvariable=self.status).grid(row=3, column=0, sticky="w", pady=(10, 0))

    def refresh_catalog(self) -> None:
        if self.downloading:
            return
        self.status.set("Refreshing Piper voice catalog...")
        threading.Thread(target=self.load_catalog, daemon=True).start()

    def load_catalog(self) -> None:
        try:
            with urlopen(PIPER_VOICES_JSON_URL, timeout=15) as response:
                catalog = json.load(response)
        except Exception as exc:
            error_message = f"Failed to load Piper catalog: {exc}"
            self.window.after(0, lambda message=error_message: self.status.set(message))
            return

        self.catalog = catalog
        self.window.after(0, self.apply_filters)

    def installed_codes(self) -> set[str]:
        return {
            model_path.stem
            for model_path in PIPER_VOICE_DIR.glob("*.onnx")
            if model_path.with_suffix(".onnx.json").exists()
        } if PIPER_VOICE_DIR.exists() else set()

    def apply_filters(self, *_args) -> None:
        installed_codes = self.installed_codes()
        search_text = self.search.get().strip().lower()
        quality = self.quality.get().strip()

        for item in self.tree.get_children():
            self.tree.delete(item)

        self.filtered_codes = []
        for voice_code in sorted(self.catalog):
            info = self.catalog[voice_code]
            label = label_for_piper_voice(voice_code, info)
            haystack = f"{label} {voice_code}".lower()
            is_installed = voice_code in installed_codes

            if search_text and search_text not in haystack:
                continue
            if quality != "all" and info.get("quality") != quality:
                continue
            if self.installed_only.get() and not is_installed:
                continue

            self.filtered_codes.append(voice_code)
            self.tree.insert(
                "",
                "end",
                iid=voice_code,
                values=(label, voice_code, info.get("quality", ""), "Yes" if is_installed else "No"),
            )

        self.status.set(f"{len(self.filtered_codes)} voice(s) shown.")

    def selected_voice_code(self) -> str | None:
        selection = self.tree.selection()
        if not selection:
            messagebox.showinfo("No selection", "Select a Piper voice first.")
            return None
        return selection[0]

    def download_selected(self) -> None:
        voice_code = self.selected_voice_code()
        if not voice_code or self.downloading:
            return

        self.downloading = True
        self.status.set(f"Downloading {voice_code}...")
        threading.Thread(target=self._download_voice_worker, args=(voice_code,), daemon=True).start()

    def _download_voice_worker(self, voice_code: str) -> None:
        try:
            import shutil

            PIPER_VOICE_DIR.mkdir(parents=True, exist_ok=True)
            info = self.catalog.get(voice_code, {})
            files = info.get("files", {})

            for rel_path in files:
                if not rel_path.endswith((".onnx", ".onnx.json")):
                    continue

                local_path = PIPER_VOICE_DIR / Path(rel_path).name
                if local_path.exists() and local_path.stat().st_size > 0:
                    continue

                url = f"https://huggingface.co/rhasspy/piper-voices/resolve/main/{quote(rel_path)}?download=true"
                with urlopen(url, timeout=60) as resp:
                    with open(local_path, "wb") as f:
                        shutil.copyfileobj(resp, f)
        except Exception as exc:
            error_message = f"Download failed: {exc}"
            self.window.after(0, lambda message=error_message: self.status.set(message))
        else:
            def on_done() -> None:
                self.status.set(f"Downloaded {voice_code}.")
                self.app.reload_piper_voices(preferred_code=voice_code)
                self.apply_filters()
            self.window.after(0, on_done)
        finally:
            self.downloading = False

    def set_selected_default(self) -> None:
        voice_code = self.selected_voice_code()
        if not voice_code:
            return

        label = self.app.find_piper_label_by_code(voice_code)
        if label is None:
            messagebox.showinfo("Voice not available", "Download the voice first, then set it as default.")
            return

        self.app.set_default_piper_voice(label)
        self.status.set(f"Default Piper voice set to {voice_code}.")

    def _rebuild_styles(self) -> None:
        """Refresh wizard styles after theme change."""
        self.tree.configure(style="Treeview")
        self.status_label = getattr(self, 'status_label', None)
        if self.status_label:
            self.status_label.configure(style="TLabel")
        self.window.update_idletasks()


class DocumentToAudioWizard:
    MP3_PRESETS = MP3_QUALITY_PRESETS
    OGG_PRESETS = OGG_QUALITY_PRESETS
    MAX_MERGE_CHUNKS = MAX_MERGE_CHUNKS

    def __init__(self, app: "App") -> None:
        self.app = app
        self.window = Toplevel(app.root)
        self.window.transient(app.root)
        self.window.title("Document to Audio Converter")
        self.window.geometry("1040x900")
        self.window.minsize(940, 780)

        self.documents: list[Path] = []
        self.doc_status: dict[Path, str] = {}
        self.extracted_texts: dict[ChapterEntry, str] = {}
        self.chunk_counts: dict[ChapterEntry, int] = {}
        self._chapter_entries: dict[Path, list[ChapterEntry]] = {}
        self._total_chunks: int = 0

        self.worker: threading.Thread | None = None
        self.pause_event = threading.Event()
        self.pause_event.set()
        self.stop_event = threading.Event()

        self.output_format = StringVar(value="MP3")
        default_quality = list(self.MP3_PRESETS.keys())[0]
        self.quality_preset = StringVar(value=default_quality)
        self.merge_files = BooleanVar(value=False)
        self.output_folder = StringVar(value=str(get_default_music_folder().resolve()))
        self.split_chapters = BooleanVar(value=False)
        self.chapter_level = StringVar(value="all")

        self.wizard_speed = DoubleVar(value=self.app.speed.get())

        self.from_page = StringVar(value="")
        self.to_page = StringVar(value="")

        initial_lang = app.language.get().strip() or "hu"
        self.wizard_language = StringVar(value=initial_lang)
        self.wizard_language_display = StringVar(value=language_display_name(initial_lang))
        self.wizard_engine = StringVar(value=app.engine.get())
        self.wizard_piper_voice_label = StringVar(value=app.piper_voice_label.get())
        self.wizard_speaker_name = StringVar(value=app.speaker_name.get())
        self.wizard_pocket_voice = StringVar(value=pocket_default_voice(initial_lang))
        self.wizard_speaker_wav = StringVar(value=app.speaker_wav.get())
        self.wizard_voice_label = StringVar(value="Voice")
        self._syncing_wizard = False

        self.phase_text = StringVar(value="")
        self.overall_text = StringVar(value="Idle")
        self.file_text = StringVar(value="")
        self.pause_button_text = StringVar(value="Pause")

        self._build_ui()
        self._on_split_chapters_toggled()
        self.output_format.trace_add("write", self._on_format_changed)
        self.wizard_language_display.trace_add("write", self._on_wizard_language_changed)
        self.wizard_engine.trace_add("write", self._on_wizard_engine_changed)
        self.wizard_speaker_wav.trace_add("write", self._on_wizard_engine_changed)
        self._on_wizard_engine_changed()
        self.window.protocol("WM_DELETE_WINDOW", self._on_close)

    def _build_ui(self) -> None:  # pragma: no cover
        frame = ttk.Frame(self.window, padding=12)
        frame.pack(fill="both", expand=True)

        header = ttk.Frame(frame)
        header.pack(fill="x", pady=(0, 8))
        header.columnconfigure(0, weight=1)

        ttk.Label(header, text="Document to Audio Converter", font=(FONT_BODY, 14, "bold")).grid(
            row=0, column=0, sticky="w"
        )

        header_actions = ttk.Frame(header)
        header_actions.grid(row=0, column=1, sticky="e")
        self.start_button = ttk.Button(
            header_actions, text="Start", style="Accent.TButton", command=self._start_processing
        )
        self.start_button.pack(side="left")
        self.pause_button = ttk.Button(
            header_actions, textvariable=self.pause_button_text, command=self._toggle_pause
        )
        self.pause_button.pack(side="left", padx=(8, 0))
        ttk.Button(header_actions, text="Stop", command=self._stop_processing).pack(
            side="left", padx=(8, 0)
        )
        ttk.Button(header_actions, text="Close", command=self._on_close).pack(
            side="left", padx=(8, 0)
        )

        # Reserve the bottom of the window for settings + progress FIRST, so the
        # progress bars are always on screen when the window opens. The document
        # list then expands into whatever space is left above them.
        bottom = ttk.Frame(frame)
        bottom.pack(side="bottom", fill="x")

        docs_frame = ttk.LabelFrame(frame, text="Documents", padding=10)
        docs_frame.pack(fill="both", expand=True)
        docs_frame.columnconfigure(0, weight=1)

        columns = ("filename", "format", "size", "status")
        self.tree = ttk.Treeview(docs_frame, columns=columns, show="headings", height=8)
        self.tree.heading("filename", text="Filename")
        self.tree.heading("format", text="Format")
        self.tree.heading("size", text="Size")
        self.tree.heading("status", text="Status")
        self.tree.column("filename", width=360)
        self.tree.column("format", width=70, anchor="center")
        self.tree.column("size", width=80, anchor="center")
        self.tree.column("status", width=280)
        self.tree.grid(row=0, column=0, sticky="nsew")

        scrollbar = ttk.Scrollbar(docs_frame, orient="vertical", command=self.tree.yview)
        scrollbar.grid(row=0, column=1, sticky="ns")
        self.tree.configure(yscrollcommand=scrollbar.set)

        doc_actions = ttk.Frame(docs_frame)
        doc_actions.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(8, 0))
        ttk.Button(doc_actions, text="Add Documents", command=self._add_documents).pack(side="left")
        ttk.Button(doc_actions, text="Remove Selected", command=self._remove_selected).pack(side="left", padx=(8, 0))
        ttk.Button(doc_actions, text="Clear All", command=self._clear_all).pack(side="right")

        settings_frame = ttk.LabelFrame(bottom, text="Output Settings", padding=10)
        settings_frame.pack(fill="x", pady=(10, 0))
        settings_frame.columnconfigure(1, weight=1)
        settings_frame.columnconfigure(3, weight=1)
        settings_frame.columnconfigure(5, weight=1)

        # Language first, then engine — same model as the main window.
        ttk.Label(settings_frame, text="Language").grid(row=0, column=0, sticky="w", pady=4)
        self.wizard_language_box = ttk.Combobox(
            settings_frame,
            textvariable=self.wizard_language_display,
            values=[language_display_name(c) for c in available_languages(self.app.piper_voice_options)],
            state="readonly",
            width=14,
        )
        self.wizard_language_box.grid(row=0, column=1, sticky="w", pady=4, padx=(0, 18))

        ttk.Label(settings_frame, text="Engine").grid(row=0, column=2, sticky="w", pady=4)
        self.wizard_engine_box = ttk.Combobox(
            settings_frame,
            textvariable=self.wizard_engine,
            values=[ENGINE_AUTO, ENGINE_PIPER, ENGINE_XTTS, ENGINE_POCKET],
            state="readonly",
            width=14,
        )
        self.wizard_engine_box.grid(row=0, column=3, sticky="w", pady=4, padx=(0, 18))

        ttk.Label(settings_frame, text="Format").grid(row=0, column=4, sticky="w", pady=4)
        self.format_box = ttk.Combobox(
            settings_frame,
            textvariable=self.output_format,
            values=list(SUPPORTED_OUTPUT_FORMATS.values()),
            state="readonly",
            width=10,
        )
        self.format_box.grid(row=0, column=5, sticky="w", pady=4)
        self.format_box.set("MP3")

        # Adaptive voice control: Piper voice list, Pocket voice list, or XTTS
        # speaker name — whichever the resolved engine uses.
        ttk.Label(settings_frame, textvariable=self.wizard_voice_label).grid(
            row=1, column=0, sticky="w", pady=4
        )
        wizard_voice_area = ttk.Frame(settings_frame)
        wizard_voice_area.grid(row=1, column=1, columnspan=3, sticky="ew", pady=4)
        wizard_voice_area.columnconfigure(0, weight=1)
        self.wizard_piper_voice_box = ttk.Combobox(
            wizard_voice_area,
            textvariable=self.wizard_piper_voice_label,
            values=list(self.app.piper_voice_options.keys()),
            state="readonly",
        )
        self.wizard_pocket_voice_box = ttk.Combobox(
            wizard_voice_area,
            textvariable=self.wizard_pocket_voice,
            values=POCKET_PREDEFINED_VOICES,
            state="readonly",
        )
        self.wizard_speaker_name_entry = ttk.Entry(
            wizard_voice_area, textvariable=self.wizard_speaker_name
        )
        for widget in (
            self.wizard_piper_voice_box,
            self.wizard_pocket_voice_box,
            self.wizard_speaker_name_entry,
        ):
            widget.grid(row=0, column=0, sticky="ew")
            widget.grid_remove()

        ttk.Label(settings_frame, text="Quality").grid(row=1, column=4, sticky="w", pady=4)
        self.quality_box = ttk.Combobox(
            settings_frame,
            textvariable=self.quality_preset,
            values=list(self.MP3_PRESETS.keys()),
            state="readonly",
            width=22,
        )
        self.quality_box.grid(row=1, column=5, sticky="w", pady=4)

        ttk.Label(settings_frame, text="Reference WAV").grid(row=2, column=0, sticky="w", pady=4)
        self.wizard_speaker_wav_entry = ttk.Entry(
            settings_frame, textvariable=self.wizard_speaker_wav
        )
        self.wizard_speaker_wav_entry.grid(row=2, column=1, columnspan=4, sticky="ew", pady=4)
        self.wizard_speaker_wav_button = ttk.Button(
            settings_frame, text="Browse", command=self._pick_wizard_reference_wav
        )
        self.wizard_speaker_wav_button.grid(row=2, column=5, sticky="e", pady=4)

        self.wizard_speed_label = ttk.Label(settings_frame, text="Speed: 1.0x")
        self.wizard_speed_label.grid(row=3, column=0, sticky="w", pady=4)
        self.wizard_speed_slider = ttk.Scale(
            settings_frame,
            from_=0.5,
            to=2.0,
            variable=self.wizard_speed,
            orient="horizontal",
            command=self._on_wizard_speed_changed,
        )
        self.wizard_speed_slider.grid(row=3, column=1, columnspan=5, sticky="ew", pady=4)

        ttk.Checkbutton(
            settings_frame,
            text="Merge all documents into one audio file",
            variable=self.merge_files,
        ).grid(row=4, column=0, columnspan=6, sticky="w", pady=(6, 4))

        split_row = ttk.Frame(settings_frame)
        split_row.grid(row=5, column=0, columnspan=6, sticky="ew", pady=(6, 4))
        self.split_chapters_check = ttk.Checkbutton(
            split_row,
            text="Split by chapter (experimental)",
            variable=self.split_chapters,
            command=self._on_split_chapters_toggled,
        )
        self.split_chapters_check.pack(side="left")
        ttk.Label(split_row, text="Level:").pack(side="left", padx=(12, 4))
        self.chapter_level_box = ttk.Combobox(
            split_row,
            textvariable=self.chapter_level,
            values=["all", "h1", "h1-h2", "h1-h3"],
            state="readonly",
            width=8,
        )
        self.chapter_level_box.pack(side="left")

        ttk.Label(settings_frame, text="Output Folder").grid(row=6, column=0, sticky="w", pady=4)
        ttk.Entry(settings_frame, textvariable=self.output_folder).grid(
            row=6, column=1, columnspan=4, sticky="ew", pady=4
        )
        ttk.Button(settings_frame, text="Browse...", command=self._pick_output_folder).grid(
            row=6, column=5, sticky="e", pady=4, padx=(8, 0)
        )

        ttk.Label(settings_frame, text="Pages (PDF)").grid(row=7, column=0, sticky="w", pady=4)
        page_row = ttk.Frame(settings_frame)
        page_row.grid(row=7, column=1, columnspan=5, sticky="w", pady=4)
        ttk.Label(page_row, text="From:").pack(side="left")
        self.from_page_entry = ttk.Entry(page_row, textvariable=self.from_page, width=6)
        self.from_page_entry.pack(side="left", padx=(4, 12))
        ttk.Label(page_row, text="To:").pack(side="left")
        self.to_page_entry = ttk.Entry(page_row, textvariable=self.to_page, width=6)
        self.to_page_entry.pack(side="left", padx=(4, 0))

        progress_frame = ttk.LabelFrame(bottom, text="Progress", padding=10)
        progress_frame.pack(fill="x", pady=(10, 0))
        progress_frame.columnconfigure(0, weight=1)

        ttk.Label(progress_frame, textvariable=self.phase_text, style="Hint.TLabel").grid(
            row=0, column=0, columnspan=2, sticky="w"
        )

        ttk.Label(progress_frame, text="Overall:").grid(row=1, column=0, sticky="w", pady=(6, 0))
        ttk.Label(progress_frame, textvariable=self.overall_text, style="Hint.TLabel").grid(
            row=1, column=1, sticky="e", pady=(6, 0)
        )
        self.overall_bar = ttk.Progressbar(progress_frame, orient="horizontal", mode="determinate", maximum=100)
        self.overall_bar.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(2, 6))

        ttk.Label(progress_frame, text="File:").grid(row=3, column=0, sticky="w")
        ttk.Label(progress_frame, textvariable=self.file_text, style="Hint.TLabel").grid(
            row=3, column=1, sticky="e"
        )
        self.file_bar = ttk.Progressbar(progress_frame, orient="horizontal", mode="determinate", maximum=100)
        self.file_bar.grid(row=4, column=0, columnspan=2, sticky="ew", pady=(2, 0))

    def _on_wizard_speed_changed(self, *args) -> None:
        speed_val = round(self.wizard_speed.get(), 1)
        self.wizard_speed_label.configure(text=f"Speed: {speed_val:.1f}x")

    def _parse_page_range(self) -> tuple[int | None, int | None]:
        from_str = self.from_page.get().strip()
        to_str = self.to_page.get().strip()
        from_page = int(from_str) if from_str else None
        to_page = int(to_str) if to_str else None
        return from_page, to_page

    def _resolved_wizard_engine(self) -> str:
        return select_engine(
            self.wizard_engine.get(),
            self.wizard_language.get(),
            bool(self.wizard_speaker_wav.get().strip()),
            self.app.piper_voice_options,
        )

    def _on_wizard_language_changed(self, *_args) -> None:
        self.wizard_language.set(
            language_code_from_display(self.wizard_language_display.get())
        )
        self._on_wizard_engine_changed()

    def _on_wizard_engine_changed(self, *_args) -> None:
        if getattr(self, "_syncing_wizard", False):
            return
        self._syncing_wizard = True
        try:
            self._sync_wizard_voice_settings()
        finally:
            self._syncing_wizard = False

    def _sync_wizard_voice_settings(self) -> None:
        language = self.wizard_language.get()

        supported = engines_supporting_language(language, self.app.piper_voice_options)
        engine_values = [ENGINE_AUTO] + supported
        if hasattr(self, "wizard_engine_box"):
            self.wizard_engine_box.configure(values=engine_values)
        if self.wizard_engine.get() not in engine_values:
            self.app.enqueue_log(
                f"{self.wizard_engine.get()} can't speak {language_display_name(language)} "
                "— switched to Auto."
            )
            self.wizard_engine.set(ENGINE_AUTO)

        resolved = self._resolved_wizard_engine()

        if resolved == ENGINE_PIPER and hasattr(self, "wizard_piper_voice_box"):
            labels = piper_voices_for_language(self.app.piper_voice_options, language) or list(
                self.app.piper_voice_options.keys()
            )
            self.wizard_piper_voice_box.configure(values=labels)
            if self.wizard_piper_voice_label.get() not in labels and labels:
                self.wizard_piper_voice_label.set(labels[0])

        if resolved == ENGINE_POCKET and is_pocket_default_voice(self.wizard_pocket_voice.get().strip()):
            self.wizard_pocket_voice.set(pocket_default_voice(language))

        self._show_wizard_voice_widget(resolved)

        clone_applicable = self.wizard_engine.get() == ENGINE_AUTO or resolved in (ENGINE_XTTS, ENGINE_POCKET)
        clone_state = "normal" if clone_applicable else "disabled"
        if hasattr(self, "wizard_speaker_wav_entry"):
            self.wizard_speaker_wav_entry.configure(state=clone_state)
            self.wizard_speaker_wav_button.configure(state=clone_state)

    def _show_wizard_voice_widget(self, resolved: str) -> None:
        specs = {
            ENGINE_PIPER: (getattr(self, "wizard_piper_voice_box", None), "Piper voice", "readonly"),
            ENGINE_POCKET: (getattr(self, "wizard_pocket_voice_box", None), "Pocket voice", "readonly"),
            ENGINE_XTTS: (getattr(self, "wizard_speaker_name_entry", None), "Built-in speaker", "normal"),
        }
        active = specs.get(resolved, specs[ENGINE_PIPER])
        if hasattr(self, "wizard_voice_label"):
            self.wizard_voice_label.set(active[1])
        for engine_key, (widget, _label, active_state) in specs.items():
            if widget is None:
                continue
            if engine_key == resolved:
                widget.configure(state=active_state)
                widget.grid()
            else:
                widget.configure(state="disabled")
                widget.grid_remove()

    def _on_format_changed(self, *_args) -> None:
        fmt = self.output_format.get()
        if fmt == "OGG":
            presets = list(self.OGG_PRESETS.keys())
            self.quality_preset.set(presets[0])
            self.quality_box.configure(values=presets, state="readonly")
        elif fmt == "WAV":
            self.quality_box.configure(values=["Lossless (no quality setting)"], state="disabled")
            self.quality_preset.set("Lossless (no quality setting)")
        else:
            presets = list(self.MP3_PRESETS.keys())
            self.quality_preset.set(presets[0])
            self.quality_box.configure(values=presets, state="readonly")

    def _on_split_chapters_toggled(self) -> None:
        enabled = self.split_chapters.get()
        if hasattr(self, "chapter_level_box"):
            self.chapter_level_box.configure(state="readonly" if enabled else "disabled")
        if not hasattr(self, "start_button") or not hasattr(self, "tree"):
            return
        if self.extracted_texts:
            self.extracted_texts.clear()
            self.chunk_counts.clear()
            self._chapter_entries.clear()
            self._total_chunks = 0
            for doc_path in self.documents:
                self._update_doc_status(doc_path, "Ready — needs re-extraction")
            self.start_button.configure(state="disabled")
            self.app.enqueue_log("Chapter split setting changed — re-extraction required.")

    def _add_documents(self) -> None:
        formats = " ".join(f"*{ext}" for ext in DocumentExtractor.SUPPORTED)
        paths = filedialog.askopenfilenames(
            parent=self.window,
            title="Select documents",
            filetypes=[
                ("All supported documents", formats),
                ("DOCX files", "*.docx"),
                ("ODT files", "*.odt"),
                ("PDF files", "*.pdf"),
                ("EPUB files", "*.epub"),
                ("MOBI files", "*.mobi"),
                ("All files", "*.*"),
            ],
        )
        for path_str in paths:
            path = Path(path_str)
            if path not in self.documents:
                self.documents.append(path)
                self.doc_status[path] = "Ready"
        self._refresh_tree()

    def _remove_selected(self) -> None:
        selection = self.tree.selection()
        for iid in selection:
            path = Path(iid)
            if path in self.documents:
                self.documents.remove(path)
                self.doc_status.pop(path, None)
        self._refresh_tree()

    def _clear_all(self) -> None:
        self.documents.clear()
        self.doc_status.clear()
        self._refresh_tree()

    def _pick_output_folder(self) -> None:
        folder = filedialog.askdirectory(
            parent=self.window,
            title="Choose output folder",
            initialdir=self.output_folder.get(),
        )
        if folder:
            self.output_folder.set(folder)

    def _pick_wizard_reference_wav(self) -> None:
        path = filedialog.askopenfilename(
            parent=self.window,
            title="Choose reference voice WAV",
            filetypes=[("Audio files", "*.wav *.mp3 *.m4a *.flac"), ("All files", "*.*")],
        )
        if path:
            self.wizard_speaker_wav.set(path)

    def _refresh_tree(self) -> None:
        for item in self.tree.get_children():
            self.tree.delete(item)
        for path in self.documents:
            size_kb = path.stat().st_size / 1024 if path.exists() else 0
            self.tree.insert(
                "",
                "end",
                iid=str(path),
                values=(
                    path.name,
                    DocumentExtractor.SUPPORTED.get(path.suffix.lower(), "Unknown"),
                    f"{size_kb:.0f} KB" if size_kb < 1024 else f"{size_kb / 1024:.1f} MB",
                    self.doc_status.get(path, "Ready"),
                ),
            )
        self.start_button.configure(state="normal" if self.documents else "disabled")

    def _set_buttons_state(self, processing: bool) -> None:
        state = "disabled" if processing else "normal"
        self.start_button.configure(state=state)
        self.pause_button_text.set("Pause")

    def _start_processing(self) -> None:
        if self.worker and self.worker.is_alive():
            return
        if not self.documents:
            messagebox.showinfo("No documents", "Add at least one document first.")
            return

        output_folder = Path(self.output_folder.get())
        output_folder.mkdir(parents=True, exist_ok=True)

        self.doc_status = {p: "Queued" for p in self.documents}
        self.extracted_texts.clear()
        self.chunk_counts.clear()
        self._chapter_entries.clear()
        self._refresh_tree()

        self._set_buttons_state(processing=True)
        self.stop_event.clear()
        self.pause_event.set()

        self.worker = threading.Thread(target=self._run_processing, args=(output_folder,), daemon=True)
        self.worker.start()

    def _run_processing(self, output_folder: Path) -> None:
        try:
            self._do_extraction()
            if self.stop_event.is_set():
                return self._finish("Stopped.")

            self._do_preparation()
            if self.stop_event.is_set():
                return self._finish("Stopped.")

            if self.merge_files.get():
                self._do_synthesis_merged(output_folder)
            else:
                self._do_synthesis_per_file(output_folder)
        except Exception as exc:
            self.app.enqueue_log(f"Document wizard error: {exc}")
            self.window.after(0, lambda e=str(exc): self._finish(f"Error: {e}"))

    def _do_extraction(self) -> None:
        docs = list(self.documents)
        total = len(docs)
        chapter_formats = {".docx", ".odt", ".pdf", ".epub"}
        from_page, to_page = self._parse_page_range()
        for i, doc_path in enumerate(docs):
            if self.stop_event.is_set():
                return
            self.window.after(0, lambda p=doc_path, s="Extracting...": self._update_doc_status(p, s))
            self.window.after(
                0,
                lambda cur=i + 1, tot=total: self.phase_text.set(
                    f"Extracting text ({cur}/{tot})..."
                ),
            )
            self.window.after(
                0,
                lambda cur=i + 1, tot=total: self._set_overall(
                    round(cur / tot * 100) if tot > 0 else 0,
                    f"Extracting ({cur}/{tot})",
                ),
            )
            try:
                if self.split_chapters.get() and doc_path.suffix.lower() in chapter_formats:
                    raw_chapters = DocumentExtractor.extract_chapters(
                        doc_path, self.chapter_level.get(),
                        from_page=from_page, to_page=to_page,
                    )
                    entries = []
                    for idx, (title, content) in enumerate(raw_chapters):
                        stripped = content.strip()
                        if not stripped:
                            continue
                        entry = ChapterEntry(
                            source_path=doc_path,
                            index=idx,
                            title=title,
                            content=stripped,
                            word_count=len(stripped.split()),
                        )
                        entries.append(entry)

                    entries = self._merge_short_chapters(entries)
                    for entry in entries:
                        self.extracted_texts[entry] = entry.content
                    self._chapter_entries[doc_path] = entries

                    if entries:
                        titles = [
                            f"{e.title[:40] if e.title else f'Ch{e.index+1}'}"
                            for e in entries
                        ]
                        status = f"{len(entries)} chapters: {', '.join(titles)}"
                    else:
                        status = "No chapters found"
                    self.window.after(
                        0,
                        lambda p=doc_path, s=status: self._update_doc_status(p, s),
                    )
                else:
                    text = DocumentExtractor.extract_text(doc_path, from_page=from_page, to_page=to_page)
                    entry = ChapterEntry(
                        source_path=doc_path,
                        index=0,
                        title="",
                        content=text,
                        word_count=len(text.split()),
                    )
                    self.extracted_texts[entry] = text
                    self._chapter_entries[doc_path] = [entry]
                    word_count = len(text.split())
                    self.window.after(
                        0,
                        lambda p=doc_path, wc=word_count: self._update_doc_status(
                            p, f"Extracted ({wc:,} words)"
                        ),
                    )
            except Exception as exc:
                self.window.after(
                    0,
                    lambda p=doc_path, e=str(exc): self._update_doc_status(p, f"Failed: {e}"),
                )

    def _do_preparation(self) -> None:
        entries = list(self.extracted_texts.keys())
        if not entries:
            raise RuntimeError("No documents could be extracted.")

        self.window.after(0, lambda: self.phase_text.set("Preparing text chunks..."))
        self.window.after(0, lambda: self._set_overall(0, "Preparing..."))

        self.chunk_counts: dict[ChapterEntry, int] = {}
        total_chunks = 0
        for i, entry in enumerate(entries):
            if self.stop_event.is_set():
                return
            text = self.extracted_texts[entry]
            count = len(chunk_text_with_offsets(text, MAX_CHARS_PER_CHUNK))
            self.chunk_counts[entry] = max(count, 1)
            total_chunks += self.chunk_counts[entry]
            self.window.after(
                0,
                lambda cur=i + 1, tot=len(entries): self._set_overall(
                    round(cur / tot * 10),
                    f"Preparing ({cur}/{tot})",
                ),
            )

        self._total_chunks = max(total_chunks, 1)

    def _build_request(self, text: str, output_path: Path) -> SynthesisRequest:
        piper_label = self.wizard_piper_voice_label.get().strip() or DEFAULT_PIPER_VOICE_LABEL
        voice_metadata = self.app.piper_voice_options.get(
            piper_label, get_piper_voice_metadata(piper_label)
        )
        language = self.wizard_language.get().strip() or "hu"
        if self._resolved_wizard_engine() == ENGINE_POCKET:
            speaker_name = self.wizard_pocket_voice.get().strip() or pocket_default_voice(language)
        else:
            speaker_name = self.wizard_speaker_name.get().strip() or DEFAULT_SPEAKER
        return SynthesisRequest(
            text=text,
            language=language,
            output_file=output_path,
            engine=self.wizard_engine.get().strip() or ENGINE_AUTO,
            piper_voice_label=piper_label,
            piper_voice_code=voice_metadata["code"],
            speaker_name=speaker_name,
            speaker_wav=self.wizard_speaker_wav.get().strip(),
            speed=round(self.wizard_speed.get(), 1),
        )

    def _do_synthesis_per_file(self, output_folder: Path) -> None:
        entries = list(self.extracted_texts.keys())
        if not entries:
            raise RuntimeError("No documents with valid extracted text.")

        total_entries = len(entries)
        cumul_chunks = 0
        suffix = f".{self.output_format.get().lower()}"
        split_active = self.split_chapters.get()
        seen_docs: set[Path] = set()

        for i, entry in enumerate(entries):
            if self.stop_event.is_set():
                return

            source_path = entry.source_path
            if split_active and source_path not in seen_docs:
                seen_docs.add(source_path)
                all_entries = self._chapter_entries.get(source_path, [entry])
                self.window.after(
                    0,
                    lambda p=source_path: self._update_doc_status(p, "Synthesizing..."),
                )

            if split_active:
                chapter_label = f"ch {entry.index + 1}/{len(self._chapter_entries.get(source_path, [entry]))}"
                self.window.after(
                    0,
                    lambda cur=i + 1, tot=total_entries, lbl=chapter_label: self.phase_text.set(
                        f"Synthesizing {lbl} ({cur}/{tot})..."
                    ),
                )
            else:
                self.window.after(
                    0,
                    lambda cur=i + 1, tot=total_entries: self.phase_text.set(
                        f"Synthesizing document {cur}/{tot}..."
                    ),
                )

            output_path = self._chapter_output_path(entry, output_folder, suffix, split_active)

            text = self.extracted_texts[entry]
            expected = self.chunk_counts.get(entry, 1)
            self._synthesize_text(
                text,
                output_path,
                expected_chunks=expected,
                doc_index=i,
                total_docs=total_entries,
                cumul_chunks_start=cumul_chunks,
            )
            cumul_chunks += expected

            if self.stop_event.is_set():
                self.window.after(0, lambda p=source_path: self._update_doc_status(p, "Stopped"))
                return

            done_status = "Done"
            if split_active:
                all_entries = self._chapter_entries.get(source_path, [])
                last_in_doc = (entry.index == all_entries[-1].index) if all_entries else True
                if not last_in_doc:
                    continue
            self.window.after(0, lambda p=source_path, s=done_status: self._update_doc_status(p, s))

        self._finish(None)

    def _do_synthesis_merged(self, output_folder: Path) -> None:
        entries = list(self.extracted_texts.keys())
        if not entries:
            raise RuntimeError("No documents with valid extracted text.")

        parts: list[str] = []
        for entry in entries:
            if entry.title:
                parts.append(entry.title)
                parts.append(entry.content)
            else:
                parts.append(entry.content)

        merged_text = "\n\n".join(parts)
        total_expected = sum(self.chunk_counts.get(e, 1) for e in entries)

        timestamp = time.strftime("%Y%m%d_%H%M%S")
        suffix = f".{self.output_format.get().lower()}"
        output_path = output_folder / f"merged_{timestamp}{suffix}"

        self.window.after(0, lambda: self.phase_text.set("Synthesizing merged output..."))

        self._synthesize_text(
            merged_text,
            output_path,
            expected_chunks=total_expected,
            doc_index=0,
            total_docs=1,
            cumul_chunks_start=0,
        )

        if not self.stop_event.is_set():
            self._finish(None)

    def _chapter_output_path(
        self,
        entry: ChapterEntry,
        output_folder: Path,
        suffix: str,
        split_active: bool,
    ) -> Path:
        if split_active and entry.title:
            safe_title = sanitize_filename(entry.title)[:50]
            return output_folder / f"{entry.source_path.stem}__{safe_title}{suffix}"
        if split_active:
            return output_folder / f"{entry.source_path.stem}_ch{entry.index + 1:02d}{suffix}"
        return output_folder / f"{entry.source_path.stem}{suffix}"

    @staticmethod
    def _merge_short_chapters(
        entries: list[ChapterEntry],
        min_words: int = MIN_CHAPTER_WORDS,
    ) -> list[ChapterEntry]:
        if not entries:
            return entries

        merged: list[ChapterEntry] = []
        i = 0
        while i < len(entries):
            current = entries[i]
            if current.word_count >= min_words or len(merged) == 0:
                merged.append(current)
                i += 1
                continue

            prev = merged[-1]
            combined_title = f"{prev.title}; {current.title}" if current.title else prev.title
            combined_content = f"{prev.content}\n\n{current.content}"
            merged[-1] = ChapterEntry(
                source_path=prev.source_path,
                index=prev.index,
                title=combined_title,
                content=combined_content,
                word_count=prev.word_count + current.word_count,
            )
            i += 1

        if merged:
            for idx, entry in enumerate(merged):
                merged[idx] = ChapterEntry(
                    source_path=entry.source_path,
                    index=idx,
                    title=entry.title,
                    content=entry.content,
                    word_count=entry.word_count,
                )

        return merged

    def _synthesize_text(
        self,
        text: str,
        output_path: Path,
        expected_chunks: int,
        doc_index: int,
        total_docs: int,
        cumul_chunks_start: int,
    ) -> None:
        request = self._build_request(text, output_path)
        combined = AudioSegment.silent(duration=0)
        chunk_count = 0

        for _chunk, segment in self.app.service.iter_segments(request):
            if self.stop_event.is_set():
                return
            self.pause_event.wait()

            combined += segment
            chunk_count += 1
            if chunk_count < expected_chunks:
                combined += AudioSegment.silent(duration=PAUSE_MS)

            current = cumul_chunks_start + chunk_count
            total = max(self._total_chunks, 1)
            overall_pct = round(current / total * 100)
            file_pct = (
                round(chunk_count / expected_chunks * 100)
                if expected_chunks > 0
                else 0
            )
            self.window.after(
                0,
                lambda ov=overall_pct, msg=f"{current}/{total} chunks": self._set_overall(ov, msg),
            )
            self.window.after(
                0,
                lambda fp=file_pct, msg=f"Chunk {chunk_count}/{expected_chunks}": self._set_file(fp, msg),
            )
            self.window.after(
                0,
                lambda cur=doc_index + 1, tot=total_docs: self.phase_text.set(
                    f"Synthesizing document {cur}/{tot}..."
                ),
            )

        if self.stop_event.is_set():
            return

        quality = self.quality_preset.get()
        preset: dict = {}
        if self.output_format.get() == "OGG":
            preset = self.OGG_PRESETS.get(quality, {})
        elif self.output_format.get() == "MP3":
            preset = self.MP3_PRESETS.get(quality, {})

        self.window.after(0, lambda: self.phase_text.set("Exporting audio file..."))
        export_audio_segment(
            combined,
            output_path,
            bitrate=preset.get("bitrate"),
            quality_params=preset.get("quality_params"),
        )
        self.app.enqueue_log(f"Exported: {output_path}")

    def _toggle_pause(self) -> None:
        if self.pause_event.is_set():
            self.pause_event.clear()
            self.pause_button_text.set("Resume")
        else:
            self.pause_event.set()
            self.pause_button_text.set("Pause")

    def _stop_processing(self) -> None:
        self.stop_event.set()
        self.pause_event.set()
        self.pause_button_text.set("Pause")

    def _on_close(self) -> None:
        if self.worker and self.worker.is_alive():
            if not messagebox.askyesno(
                "Processing in progress",
                "A conversion job is running. Stop it and close?",
            ):
                return
            self._stop_processing()
            self.worker.join(timeout=3)
        self.window.destroy()

    def _update_doc_status(self, doc_path: Path, status: str) -> None:
        self.doc_status[doc_path] = status
        iid = str(doc_path)
        if self.tree.exists(iid):
            values = list(self.tree.item(iid, "values"))
            if len(values) >= 4:
                values[3] = status
                self.tree.item(iid, values=values)

    def _set_overall(self, value: int, text: str) -> None:
        self.overall_bar["value"] = value
        self.overall_text.set(text)

    def _set_file(self, value: int, text: str) -> None:
        self.file_bar["value"] = value
        self.file_text.set(text)

    def _finish(self, error: str | None) -> None:
        def apply() -> None:
            self._set_buttons_state(processing=False)
            if error:
                self.phase_text.set(error)
                self.overall_text.set("Failed")
            else:
                self.phase_text.set("Conversion complete.")
                self.overall_bar["value"] = 100
                self.overall_text.set("Done")
                self.file_bar["value"] = 100
                self.file_text.set("")
                self.app.enqueue_log("Document conversion complete.")
        self.window.after(0, apply)

    def _rebuild_styles(self) -> None:
        """Refresh wizard styles after theme change."""
        self.tree.configure(style="Treeview")
        self.overall_bar.configure(style="TProgressbar")
        self.file_bar.configure(style="TProgressbar")
        self.window.update_idletasks()


class LanguageLearningWizard:
    """Wizard for generating and practicing language learning sentence pairs."""

    def __init__(self, app: "App") -> None:
        self.app = app
        self.window = Toplevel(app.root)
        self.window.transient(app.root)
        self.window.title("Language Learning Practice")
        self.window.geometry("900x700")
        self.window.minsize(800, 600)

        # Load only normalized, canonical language settings.
        self.settings = normalize_app_settings(app.settings)["language_learning"]
        self.availability = language_learning_availability()

        # State variables
        self.preset_var = StringVar(value=self.settings.get("preset", "A2"))
        self.lang_var = StringVar(value=language_display_name(normalize_learning_language(self.settings.get("language"))))
        self.level_var = IntVar(value=self.settings.get("level", 1))
        self.count_var = IntVar(value=self.settings.get("count", 15))
        self.max_length_var = IntVar(value=self.settings.get("max_length", 80))
        self.plural_chance_var = DoubleVar(value=self.settings.get("plural_chance", 0.3))
        self.seed_var = StringVar(value=str(self.settings.get("seed") or ""))
        self.top_n_var = StringVar(value=str(self.settings.get("top_n") or ""))
        self.base_word_var = StringVar(value=self.settings.get("base_word") or "")
        self.base_word_count_var = IntVar(value=self.settings.get("base_word_count", 10))
        self.base_template_var = StringVar(value=self.settings.get("base_template") or "")
        self.vary_role_var = StringVar(value=self.settings.get("vary_role") or "")
        self.vary_words_var = StringVar(value=self.settings.get("vary_words") or "")

        # TTS settings
        self.engine_var = StringVar(value=app.engine.get())
        self.pair_pause_var = IntVar(value=self.settings.get("pair_pause_ms", 2000))
        self.auto_speak_var = BooleanVar(value=self.settings.get("auto_speak", False))
        self.show_trans_var = BooleanVar(value=self.settings.get("show_translations", True))
        self.speed_var = DoubleVar(value=app.speed.get())

        # Voice selections (synced from main window)
        self.piper_voice_label_var = StringVar(value=app.piper_voice_label.get())
        self.pocket_voice_var = StringVar(value=app.pocket_voice.get())
        self.speaker_name_var = StringVar(value=app.speaker_name.get())
        self.speaker_wav_var = StringVar(value=app.speaker_wav.get())

        # Generated content
        self.generated_pairs: list[tuple[str, str]] = []

        # Worker thread
        self.worker: threading.Thread | None = None
        self.stop_event = threading.Event()
        self.pause_event = threading.Event()
        self.pause_event.set()
        self.session_state = SessionState.IDLE
        self._job_id = 0

        # Temp files for cleanup
        self._temp_files: list[Path] = []

        self._build_ui()
        if not self.availability.available:
            self.window.protocol("WM_DELETE_WINDOW", self._on_close)
            return
        self._validate_language()
        self._apply_preset()
        self._sync_voice_settings()

        # Traces
        self.preset_var.trace_add("write", self._on_preset_changed)
        self.lang_var.trace_add("write", self._on_language_changed)
        self.engine_var.trace_add("write", self._on_engine_changed)

        self.window.protocol("WM_DELETE_WINDOW", self._on_close)

    def _check_language_practice_available(self) -> None:
        """Compatibility wrapper for callers that used the old eager check."""
        self.availability = language_learning_availability()

    def _build_unavailable_ui(self) -> None:
        main = ttk.Frame(self.window, padding=28)
        main.pack(fill="both", expand=True)
        ttk.Label(main, text="Language Learning is not ready", style="Header.TLabel").pack(anchor="w")
        ttk.Label(
            main,
            text=self.availability.message,
            style="Hint.TLabel",
            wraplength=720,
            justify="left",
        ).pack(anchor="w", pady=(12, 20))
        actions = ttk.Frame(main)
        actions.pack(anchor="w")
        ttk.Button(actions, text="Retry", style="Accent.TButton", command=self._retry_availability).pack(side="left")
        ttk.Button(actions, text="Close", command=self._on_close).pack(side="left", padx=(8, 0))

    def _retry_availability(self) -> None:
        self.availability = language_learning_availability()
        if self.availability.available:
            self.window.destroy()
            self.app.lang_learning_wizard = LanguageLearningWizard(self.app)
            return
        for child in self.window.winfo_children():
            child.destroy()
        self._build_unavailable_ui()

    def _build_ui(self) -> None:
        if not self.availability.available:
            self._build_unavailable_ui()
            return
        main = ttk.Frame(self.window, padding=12)
        main.pack(fill="both", expand=True)
        main.columnconfigure(0, weight=1)
        main.rowconfigure(3, weight=1)

        # Header with preset and language
        header = ttk.Frame(main)
        header.grid(row=0, column=0, sticky="ew", pady=(0, 12))
        header.columnconfigure(2, weight=1)

        ttk.Label(header, text="Preset").grid(row=0, column=0, sticky="w", padx=(0, 6))
        self.preset_box = ttk.Combobox(
            header,
            textvariable=self.preset_var,
            values=["A1", "A2", "B1", "B2", "Custom"],
            state="readonly",
            width=8,
        )
        self.preset_box.grid(row=0, column=1, sticky="w", padx=(0, 18))

        ttk.Label(header, text="Language").grid(row=0, column=2, sticky="w", padx=(0, 6))
        self.lang_box = ttk.Combobox(
            header,
            textvariable=self.lang_var,
            values=["Portuguese", "Spanish"],
            state="readonly",
            width=12,
        )
        self.lang_box.grid(row=0, column=3, sticky="w", padx=(0, 18))

        ttk.Button(header, text="Save as Defaults", command=self._save_preset).grid(row=0, column=4, sticky="e")

        # Settings panels
        settings_panels = ttk.Frame(main)
        settings_panels.grid(row=1, column=0, sticky="ew", pady=(0, 12))
        settings_panels.columnconfigure(0, weight=1)
        settings_panels.columnconfigure(1, weight=1)

        # Generator Settings
        gen_frame = ttk.LabelFrame(settings_panels, text="Generator Settings", padding=10)
        gen_frame.grid(row=0, column=0, sticky="nsew", padx=(0, 6))
        gen_frame.columnconfigure(1, weight=1)

        row = 0
        ttk.Label(gen_frame, text="Level").grid(row=row, column=0, sticky="w", pady=4)
        self.level_box = ttk.Combobox(gen_frame, textvariable=self.level_var, values=[0, 1, 2, 3], state="readonly", width=5)
        self.level_box.grid(row=row, column=1, sticky="w", pady=4)
        row += 1

        ttk.Label(gen_frame, text="Count").grid(row=row, column=0, sticky="w", pady=4)
        make_validated_entry(gen_frame, self.count_var, validator="int", width=8).grid(row=row, column=1, sticky="w", pady=4)
        row += 1

        ttk.Label(gen_frame, text="Max words").grid(row=row, column=0, sticky="w", pady=4)
        make_validated_entry(gen_frame, self.max_length_var, validator="int", width=8).grid(row=row, column=1, sticky="w", pady=4)
        row += 1

        ttk.Label(gen_frame, text="Plural %").grid(row=row, column=0, sticky="w", pady=4)
        make_validated_entry(gen_frame, self.plural_chance_var, validator="float", width=8).grid(row=row, column=1, sticky="w", pady=4)
        row += 1

        ttk.Label(gen_frame, text="Seed (optional)").grid(row=row, column=0, sticky="w", pady=4)
        make_validated_entry(gen_frame, self.seed_var, validator="int", width=12).grid(row=row, column=1, sticky="w", pady=4)
        row += 1

        ttk.Label(gen_frame, text="Top N words (optional)").grid(row=row, column=0, sticky="w", pady=4)
        make_validated_entry(gen_frame, self.top_n_var, validator="int", width=12).grid(row=row, column=1, sticky="w", pady=4)
        row += 1

        # Base word drill
        ttk.Separator(gen_frame, orient="horizontal").grid(row=row, column=0, columnspan=2, sticky="ew", pady=8)
        row += 1
        ttk.Label(gen_frame, text="Base word (drill)").grid(row=row, column=0, sticky="w", pady=4)
        make_validated_entry(gen_frame, self.base_word_var, width=18).grid(row=row, column=1, sticky="w", pady=4)
        row += 1
        ttk.Label(gen_frame, text="Min count").grid(row=row, column=0, sticky="w", pady=4)
        make_validated_entry(gen_frame, self.base_word_count_var, validator="int", width=8).grid(row=row, column=1, sticky="w", pady=4)
        row += 1

        # Batch mode
        ttk.Separator(gen_frame, orient="horizontal").grid(row=row, column=0, columnspan=2, sticky="ew", pady=8)
        row += 1
        ttk.Label(gen_frame, text="Template ID").grid(row=row, column=0, sticky="w", pady=4)
        self.template_box = ttk.Combobox(gen_frame, textvariable=self.base_template_var, values=self._get_template_ids(), state="readonly", width=8)
        self.template_box.grid(row=row, column=1, sticky="w", pady=4)
        row += 1
        ttk.Label(gen_frame, text="Vary role").grid(row=row, column=0, sticky="w", pady=4)
        self.vary_role_box = ttk.Combobox(gen_frame, textvariable=self.vary_role_var, values=self._get_vary_roles(), state="readonly", width=10)
        self.vary_role_box.grid(row=row, column=1, sticky="w", pady=4)
        row += 1
        ttk.Label(gen_frame, text="Vary words (comma-separated)").grid(row=row, column=0, sticky="w", pady=4)
        make_validated_entry(gen_frame, self.vary_words_var, width=24).grid(row=row, column=1, sticky="w", pady=4)

        # TTS Settings
        tts_frame = ttk.LabelFrame(settings_panels, text="TTS Settings", padding=10)
        tts_frame.grid(row=0, column=1, sticky="nsew", padx=(6, 0))
        tts_frame.columnconfigure(1, weight=1)

        row = 0
        ttk.Label(tts_frame, text="Engine").grid(row=row, column=0, sticky="w", pady=4)
        self.engine_box = ttk.Combobox(tts_frame, textvariable=self.engine_var, values=[ENGINE_AUTO, ENGINE_PIPER, ENGINE_XTTS, ENGINE_POCKET], state="readonly", width=14)
        self.engine_box.grid(row=row, column=1, sticky="w", pady=4)
        row += 1

        ttk.Label(tts_frame, text="Voice").grid(row=row, column=0, sticky="w", pady=4)
        self.voice_area = ttk.Frame(tts_frame)
        self.voice_area.grid(row=row, column=1, sticky="ew", pady=4)
        self.voice_area.columnconfigure(0, weight=1)
        self.piper_voice_box = ttk.Combobox(self.voice_area, textvariable=self.piper_voice_label_var, values=list(self.app.piper_voice_options.keys()), state="readonly")
        self.pocket_voice_box = ttk.Combobox(self.voice_area, textvariable=self.pocket_voice_var, values=POCKET_PREDEFINED_VOICES, state="readonly")
        self.speaker_name_entry = ttk.Entry(self.voice_area, textvariable=self.speaker_name_var)
        add_context_menu(self.speaker_name_entry)
        for w in (self.piper_voice_box, self.pocket_voice_box, self.speaker_name_entry):
            w.grid(row=0, column=0, sticky="ew")
            w.grid_remove()
        row += 1

        ttk.Label(tts_frame, text="Speed").grid(row=row, column=0, sticky="w", pady=4)
        self.speed_label = ttk.Label(tts_frame, text=f"Speed: {self.speed_var.get():.1f}x")
        self.speed_label.grid(row=row, column=1, sticky="w", pady=4)
        self.speed_slider = ttk.Scale(tts_frame, from_=0.5, to=2.0, variable=self.speed_var, orient="horizontal", command=self._on_speed_changed)
        self.speed_slider.grid(row=row, column=2, sticky="ew", pady=4, padx=(8, 0))
        row += 1

        ttk.Label(tts_frame, text="Pair pause (ms)").grid(row=row, column=0, sticky="w", pady=4)
        make_validated_entry(tts_frame, self.pair_pause_var, validator="int", width=8).grid(row=row, column=1, sticky="w", pady=4)
        row += 1

        ttk.Checkbutton(tts_frame, text="Auto-speak after generate", variable=self.auto_speak_var).grid(row=row, column=0, columnspan=2, sticky="w", pady=4)
        row += 1
        ttk.Checkbutton(tts_frame, text="Show translations in output", variable=self.show_trans_var).grid(row=row, column=0, columnspan=2, sticky="w", pady=4)

        # Action buttons
        actions = ttk.Frame(main)
        actions.grid(row=2, column=0, sticky="ew", pady=(0, 12))
        self.generate_button = ttk.Button(actions, text="Generate", style="Accent.TButton", command=self._generate)
        self.generate_button.pack(side="left")
        self.speak_button = ttk.Button(actions, text="Speak", command=self._speak)
        self.speak_button.pack(side="left", padx=(8, 0))
        self.pause_button_text = StringVar(value="Pause")
        self.pause_button = ttk.Button(actions, textvariable=self.pause_button_text, command=self._toggle_pause, state="disabled")
        self.pause_button.pack(side="left", padx=(8, 0))
        self.stop_button = ttk.Button(actions, text="Stop", command=self._stop_session, state="disabled")
        self.stop_button.pack(side="left", padx=(8, 0))
        ttk.Button(actions, text="Send to Main", command=self._send_to_main).pack(side="left", padx=(8, 0))
        ttk.Button(actions, text="Export", command=self._export).pack(side="left", padx=(8, 0))
        ttk.Button(actions, text="Clear", command=self._clear).pack(side="right")

        # Output text area
        ttk.Label(main, text="Generated Sentences (editable)").grid(row=3, column=0, sticky="w", pady=(0, 4))
        text_frame = ttk.Frame(main)
        text_frame.grid(row=4, column=0, sticky="nsew")
        text_frame.columnconfigure(0, weight=1)
        text_frame.rowconfigure(0, weight=1)

        self.output_text = Text(
            text_frame,
            wrap="word",
            font=(FONT_MONO, 10),
            padx=12,
            pady=12,
            undo=True,
            bg=TEXT_BG,
            fg=THEMES[CURRENT_THEME]["label_fg"],
            relief="flat",
            insertbackground=ACCENT,
            highlightthickness=1,
            highlightbackground=TEXT_BORDER,
            highlightcolor=ACCENT,
        )
        self.output_text.grid(row=0, column=0, sticky="nsew")
        add_context_menu(self.output_text)

        scroll = ttk.Scrollbar(text_frame, orient="vertical", command=self.output_text.yview)
        scroll.grid(row=0, column=1, sticky="ns")
        self.output_text.configure(yscrollcommand=scroll.set)

        # Status bar
        self.status_var = StringVar(value="Ready")
        ttk.Label(main, textvariable=self.status_var, style="Hint.TLabel").grid(row=5, column=0, sticky="w", pady=(8, 0))

    def _get_template_ids(self) -> list[str]:
        if not self.availability.available:
            return []
        import importlib

        lang_code = normalize_learning_language(self.lang_var.get())
        lang_module = importlib.import_module(f"language_practice.languages.{lang_code}")
        return [t.id for t in lang_module.TEMPLATES]

    def _get_vary_roles(self) -> list[str]:
        return ["N_SUBJ", "N_OBJ", "N_PLACE", "N_SUBJ2", "V_intr", "V_trans", "V_loc", "V_intr2", "PRON_SUBJ", "ADV", "ADJ"]

    def _validate_language(self) -> bool:
        requested = self.lang_var.get()
        language = normalize_learning_language(requested)
        normalized_display = language_display_name(language)
        if requested != normalized_display:
            self.lang_var.set(normalized_display)
            self.status_var.set(
                f"{requested or 'Saved language'} is not supported here; switched to Portuguese."
            )
        return language in {"pt", "es"}

    def _on_preset_changed(self, *_args) -> None:
        if self.preset_var.get() != "Custom":
            self._apply_preset()

    def _apply_preset(self) -> None:
        preset = self.preset_var.get()
        if preset in CEFR_PRESETS:
            p = CEFR_PRESETS[preset]
            self.level_var.set(p["level"])
            self.count_var.set(p["count"])
            self.max_length_var.set(p["max_length"])
            self.plural_chance_var.set(p["plural_chance"])
            self.settings.update(p)

    def _on_language_changed(self, *_args) -> None:
        self._validate_language()
        self._sync_voice_settings()
        # Update template IDs for new language
        self.template_box.configure(values=self._get_template_ids())
        if self.base_template_var.get() not in self._get_template_ids():
            self.base_template_var.set("")

    def _on_engine_changed(self, *_args) -> None:
        self._sync_voice_settings()

    def _on_speed_changed(self, *_args) -> None:
        self.speed_label.configure(text=f"Speed: {self.speed_var.get():.1f}x")

    def _sync_voice_settings(self) -> None:
        """Sync voice controls based on selected engine and language."""
        engine = self.engine_var.get()
        lang_code = "pt" if self.lang_var.get() == "Portuguese" else "es"

        # Filter engines that support this language
        supported = engines_supporting_language(lang_code, self.app.piper_voice_options)
        engine_values = [ENGINE_AUTO] + supported
        self.engine_box.configure(values=engine_values)
        if engine not in engine_values:
            engine = ENGINE_AUTO
            self.engine_var.set(engine)

        # Show appropriate voice control
        for w in (self.piper_voice_box, self.pocket_voice_box, self.speaker_name_entry):
            w.grid_remove()

        if engine == ENGINE_PIPER:
            labels = piper_voices_for_language(self.app.piper_voice_options, lang_code) or list(self.app.piper_voice_options.keys())
            self.piper_voice_box.configure(values=labels)
            if self.piper_voice_label_var.get() not in labels and labels:
                self.piper_voice_label_var.set(labels[0])
            self.piper_voice_box.grid()
        elif engine == ENGINE_POCKET:
            if self.pocket_voice_var.get() not in POCKET_PREDEFINED_VOICES:
                self.pocket_voice_var.set(pocket_default_voice(lang_code))
            self.pocket_voice_box.grid()
        else:  # XTTS or Auto
            self.speaker_name_entry.grid()

    def _save_preset(self) -> None:
        """Save current settings as a custom preset."""
        self.settings = self._collect_settings()
        self.app.settings["language_learning"] = self.settings
        save_app_settings(self.app.settings)
        self.status_var.set("Preset saved")
        self.window.after(2000, lambda: self.status_var.set("Ready"))

    def _collect_settings(self) -> dict:
        return {
            "preset": self.preset_var.get(),
            "language": normalize_learning_language(self.lang_var.get()),
            "level": self.level_var.get(),
            "count": self.count_var.get(),
            "max_length": self.max_length_var.get(),
            "plural_chance": self.plural_chance_var.get(),
            "seed": int(self.seed_var.get()) if self.seed_var.get().strip() else None,
            "top_n": int(self.top_n_var.get()) if self.top_n_var.get().strip() else None,
            "base_word": self.base_word_var.get().strip() or None,
            "base_word_count": self.base_word_count_var.get(),
            "base_template": self.base_template_var.get().strip() or None,
            "vary_role": self.vary_role_var.get().strip() or None,
            "vary_words": self.vary_words_var.get().strip() or None,
            "pair_pause_ms": self.pair_pause_var.get(),
            "auto_speak": self.auto_speak_var.get(),
            "show_translations": self.show_trans_var.get(),
        }

    def _generate(self) -> None:
        """Generate sentence pairs using language-practice library."""
        if self.worker and self.worker.is_alive():
            return
        if not self.availability.available or not self._validate_language():
            return
        options = self._generation_options()
        self.stop_event.clear()
        self.pause_event.set()
        self._job_id += 1
        job_id = self._job_id
        self._set_session_state(SessionState.GENERATING)

        def worker():
            try:
                pairs = self._generate_pairs(options)
                if not self.stop_event.is_set():
                    self.window.after(0, lambda: self._display_pairs(pairs, job_id))
            except Exception as exc:
                if not self.stop_event.is_set():
                    self.window.after(0, lambda e=exc: self._on_generate_error(e))
            finally:
                if self.stop_event.is_set():
                    self.window.after(0, lambda: self._set_session_state(SessionState.IDLE, "Stopped."))

        self.worker = threading.Thread(target=worker, daemon=True)
        self.worker.start()

    def _generation_options(self) -> dict[str, object]:
        return {
            "language": normalize_learning_language(self.lang_var.get()),
            "level": self.level_var.get(),
            "count": self.count_var.get(),
            "max_length": self.max_length_var.get(),
            "plural_chance": self.plural_chance_var.get(),
            "seed": int(self.seed_var.get()) if self.seed_var.get().strip() else None,
            "top_n": int(self.top_n_var.get()) if self.top_n_var.get().strip() else None,
            "base_word": self.base_word_var.get().strip(),
            "base_word_count": self.base_word_count_var.get(),
            "base_template": self.base_template_var.get().strip(),
            "vary_role": self.vary_role_var.get().strip(),
            "vary_words": self.vary_words_var.get().strip(),
        }

    def _generate_pairs(self, options: dict[str, object]) -> list[tuple[str, str]]:
        """Call language-practice library to generate sentence pairs."""
        try:
            import importlib

            generator_cls = importlib.import_module("language_practice.generator").Generator
            lang_code = str(options["language"])
            lang_module = importlib.import_module(f"language_practice.languages.{lang_code}")
        except ImportError as exc:
            raise RuntimeError(
                "language-practice package not found. "
                "Install it with: pip install -e ../language-practice"
            ) from exc

        if self.availability.wordlist_dir is None:
            raise RuntimeError("Language-practice wordlists are unavailable. Select Retry after fixing the companion setup.")
        wordlist_path = self.availability.wordlist_dir / lang_code / "wordlist.md"
        raw_words = lang_module.parse_wordlist(wordlist_path)
        words = lang_module.enrich_words(raw_words)
        en_dict = lang_module.build_en_dict(words)

        # Create generator
        gen = generator_cls(
            words,
            level=int(options["level"]),
            max_length=int(options["max_length"]),
            plural_chance=float(options["plural_chance"]),
            seed=options["seed"],
            top_n=options["top_n"],
            lang=lang_module,
        )

        sentences: list = []

        if options["base_word"]:
            # Base word drill
            base_word = str(options["base_word"]).lower()
            # Normalize for comparison (strip accents, lower)
            import unicodedata
            def normalize(s: str) -> str:
                return "".join(c for c in unicodedata.normalize("NFD", s.lower()) if unicodedata.category(c) != "Mn")
            base_word_norm = normalize(base_word)
            base_word_obj = next((w for w in words if normalize(getattr(w, "pt", getattr(w, "es", ""))) == base_word_norm), None)
            if not base_word_obj:
                raise ValueError(f"Base word '{base_word}' not found in vocabulary")
            sentences = gen.generate_with_base_word(int(options["base_word_count"]), base_word_obj, seed=options["seed"])
        elif options["base_template"]:
            # Batch mode
            tpl = next((t for t in lang_module.TEMPLATES if t.id == options["base_template"]), None)
            if not tpl:
                raise ValueError(f"Template '{options['base_template']}' not found")
            vary_words = [w.strip() for w in str(options["vary_words"]).split(",")] if options["vary_words"] else None
            sentences = gen.generate_batch(
                int(options["count"]), tpl, str(options["vary_role"]), vary_words, seed=options["seed"]
            )
        else:
            # Normal generation
            sentences = list(gen.generate(int(options["count"])))

        # Format as pairs
        pairs = []
        for s in sentences:
            target = " ".join(s.words)
            translation = lang_module.translate(s.words, en_dict)
            pairs.append((target, translation))

        return pairs

    def _display_pairs(self, pairs: list[tuple[str, str]], job_id: int | None = None) -> None:
        if job_id is not None and job_id != self._job_id:
            return
        self.generated_pairs = pairs
        self.output_text.delete("1.0", "end")
        show_trans = self.show_trans_var.get()
        lines = []
        for i, (target, trans) in enumerate(pairs, 1):
            lines.append(f"{i}. {target}")
            if show_trans:
                lines.append(f"   {trans}")
            lines.append("")
        self.output_text.insert("1.0", "\n".join(lines))
        self._set_session_state(SessionState.IDLE, f"Generated {len(pairs)} sentence pairs")

        if self.auto_speak_var.get():
            self._speak()

    def _on_generate_error(self, exc: Exception) -> None:
        self._set_session_state(SessionState.FAILED, "Generation failed")
        messagebox.showerror("Generation Error", str(exc))

    def _speak(self) -> None:
        """Synthesize and play the generated pairs with pair pauses."""
        pairs = self._pairs_from_output()
        if not pairs:
            messagebox.showinfo("Nothing to speak", "Generate sentences first.")
            return
        if self.worker and self.worker.is_alive():
            return

        config = LanguageSessionConfig(
            language=normalize_learning_language(self.lang_var.get()),
            engine=self.engine_var.get(),
            speed=round(self.speed_var.get(), 1),
            pair_pause_ms=max(0, self.pair_pause_var.get()),
            show_translations=self.show_trans_var.get(),
        )
        self.stop_event.clear()
        self.pause_event.set()
        self._job_id += 1
        self._set_session_state(SessionState.SYNTHESIZING)

        def worker():
            try:
                combined = self._synthesize_pairs(pairs, config)
                if self.stop_event.is_set():
                    return
                import tempfile
                with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f:
                    temp_path = Path(f.name)
                self._temp_files.append(temp_path)
                export_audio_segment(combined, temp_path)
                if self.stop_event.is_set():
                    self._cleanup_temp_file(temp_path)
                    return
                self.window.after(0, lambda: self._set_session_state(SessionState.PLAYING, "Playing practice audio..."))
                self.app.player.play_blocking(temp_path, self.stop_event)
                self._cleanup_temp_file(temp_path)
                if not self.stop_event.is_set():
                    self.window.after(0, lambda: self._set_session_state(SessionState.IDLE, "Playback complete"))
            except Exception as exc:
                if not self.stop_event.is_set():
                    self.window.after(0, lambda e=exc: self._on_speak_error(e))
            finally:
                if self.stop_event.is_set():
                    self.window.after(0, lambda: self._set_session_state(SessionState.IDLE, "Stopped."))

        self.worker = threading.Thread(target=worker, daemon=True)
        self.worker.start()

    def _wait_for_session(self) -> bool:
        """Pause at a safe boundary and return False once the session is stopped."""
        while not self.pause_event.wait(timeout=0.1):
            if self.stop_event.is_set():
                return False
        return not self.stop_event.is_set()

    def _wait_pair_pause(self, milliseconds: int) -> bool:
        deadline = time.monotonic() + milliseconds / 1000
        while time.monotonic() < deadline:
            if not self._wait_for_session():
                return False
            time.sleep(min(0.05, max(0, deadline - time.monotonic())))
        return not self.stop_event.is_set()

    def _synthesize_pairs(
        self,
        pairs: list[tuple[str, str]],
        config: LanguageSessionConfig,
    ) -> AudioSegment:
        """Synthesize all pairs with pair_pause_ms between pairs."""
        combined = AudioSegment.silent(duration=0)

        for i, (target, translation) in enumerate(pairs):
            if not self._wait_for_session():
                break
            # Synthesize target
            req_target = self._build_request(target, config.language, config)
            for _chunk, segment in self.app.service.iter_segments(req_target):
                if not self._wait_for_session():
                    return combined
                combined += segment

            # Synthesize translation if enabled
            if config.show_translations and translation:
                req_trans = self._build_request(translation, "en", config)
                for _chunk, segment in self.app.service.iter_segments(req_trans):
                    if not self._wait_for_session():
                        return combined
                    combined += segment

            # Add pair pause (except after last)
            if i < len(pairs) - 1:
                if not self._wait_pair_pause(config.pair_pause_ms):
                    return combined
                combined += AudioSegment.silent(duration=config.pair_pause_ms)

        return combined

    def _build_request(
        self,
        text: str,
        language: str | None = None,
        config: LanguageSessionConfig | None = None,
    ) -> SynthesisRequest:
        """Build synthesis request using current TTS settings."""
        lang_code = normalize_learning_language(language or self.lang_var.get()) if language != "en" else "en"
        engine = config.engine if config else self.engine_var.get()
        speed = config.speed if config else round(self.speed_var.get(), 1)

        # Resolve voice
        if engine == ENGINE_PIPER:
            labels = piper_voices_for_language(self.app.piper_voice_options, lang_code)
            if not labels:
                raise RuntimeError(
                    f"Piper has no installed {language_display_name(lang_code)} voice. "
                    "Choose Auto/another engine or download a matching voice."
                )
            selected = self.piper_voice_label_var.get().strip()
            piper_label = selected if selected in labels else labels[0]
            voice_metadata = self.app.piper_voice_options[piper_label]
            piper_code = voice_metadata["code"]
            speaker_name = DEFAULT_SPEAKER
        elif engine == ENGINE_POCKET:
            piper_label = DEFAULT_PIPER_VOICE_LABEL
            voice_metadata = get_piper_voice_metadata(piper_label)
            piper_code = voice_metadata["code"]
            speaker_name = (
                self.pocket_voice_var.get().strip()
                if lang_code == normalize_learning_language(self.lang_var.get())
                else pocket_default_voice(lang_code)
            )
        else:
            labels = piper_voices_for_language(self.app.piper_voice_options, lang_code)
            piper_label = labels[0] if labels else DEFAULT_PIPER_VOICE_LABEL
            voice_metadata = get_piper_voice_metadata(piper_label)
            piper_code = voice_metadata["code"]
            speaker_name = self.speaker_name_var.get().strip() or DEFAULT_SPEAKER

        speaker_wav = self.speaker_wav_var.get().strip()

        return SynthesisRequest(
            text=text,
            language=lang_code,
            output_file=Path("temp.mp3"),
            engine=engine,
            piper_voice_label=piper_label,
            piper_voice_code=piper_code,
            speaker_name=speaker_name,
            speaker_wav=speaker_wav,
            speed=speed,
        )

    def _pairs_from_output(self) -> list[tuple[str, str]]:
        """Use edited output as the source of truth before speaking/exporting."""
        raw = self.output_text.get("1.0", "end-1c").strip()
        if not raw:
            return self.generated_pairs
        pairs: list[tuple[str, str]] = []
        for block in re.split(r"\n\s*\n", raw):
            lines = [line.strip() for line in block.splitlines() if line.strip()]
            if not lines:
                continue
            target = re.sub(r"^\d+\.\s*", "", lines[0]).strip()
            translation = lines[1] if len(lines) > 1 else ""
            if target:
                pairs.append((target, translation))
        self.generated_pairs = pairs or self.generated_pairs
        return self.generated_pairs

    def _set_session_state(self, state: SessionState, message: str | None = None) -> None:
        self.session_state = state
        if message:
            self.status_var.set(message)
        active = state in {
            SessionState.GENERATING,
            SessionState.SYNTHESIZING,
            SessionState.PLAYING,
            SessionState.PAUSED,
            SessionState.STOPPING,
        }
        can_pause = state in {SessionState.SYNTHESIZING, SessionState.PLAYING, SessionState.PAUSED}
        self.generate_button.configure(state="disabled" if active else "normal")
        self.speak_button.configure(state="disabled" if active else "normal")
        self.pause_button.configure(state="normal" if can_pause else "disabled")
        self.stop_button.configure(state="normal" if active else "disabled")
        self.pause_button_text.set("Resume" if state == SessionState.PAUSED else "Pause")

    def _toggle_pause(self) -> None:
        if self.session_state == SessionState.PAUSED:
            self.pause_event.set()
            self.app.player.resume()
            self._set_session_state(SessionState.PLAYING, "Resumed.")
            return
        if self.session_state not in {SessionState.SYNTHESIZING, SessionState.PLAYING}:
            return
        self.pause_event.clear()
        self.app.player.pause()
        self._set_session_state(SessionState.PAUSED, "Paused.")

    def _stop_session(self) -> None:
        if self.session_state == SessionState.IDLE:
            return
        self._set_session_state(SessionState.STOPPING, "Stopping...")
        self.stop_event.set()
        self.pause_event.set()
        self.app.player.stop(quiet=True)
        self._job_id += 1

    def _cleanup_temp_file(self, path: Path) -> None:
        """Remove temp file and remove from tracking list."""
        try:
            path.unlink(missing_ok=True)
        except Exception:
            pass
        if path in self._temp_files:
            self._temp_files.remove(path)

    def _on_speak_error(self, exc: Exception) -> None:
        self._set_session_state(SessionState.FAILED, "Playback failed")
        messagebox.showerror("Playback Error", str(exc))

    def _send_to_main(self) -> None:
        """Send generated text to main window."""
        pairs = self._pairs_from_output()
        if not pairs:
            return
        show_trans = self.show_trans_var.get()
        lines = []
        for target, trans in pairs:
            lines.append(target)
            if show_trans:
                lines.append(trans)
            lines.append("")
        text = "\n".join(lines)
        self.app.text.delete("1.0", "end")
        self.app.text.insert("1.0", text)
        self.app.root.lift()
        self.app.root.focus_force()
        self.status_var.set("Sent to main window")

    def _export(self) -> None:
        """Export generated pairs to file."""
        if not self._pairs_from_output():
            messagebox.showinfo("Nothing to export", "Generate sentences first.")
            return

        format_var = StringVar(value="Text")
        dialog = Toplevel(self.window)
        dialog.transient(self.window)
        dialog.title("Export Format")
        dialog.geometry("300x150")
        dialog.grab_set()

        ttk.Label(dialog, text="Choose export format:").pack(pady=12)
        fmt_box = ttk.Combobox(dialog, textvariable=format_var, values=["Text", "Anki CSV", "JSON"], state="readonly")
        fmt_box.pack(pady=6)

        def do_export():
            fmt = format_var.get()
            dialog.destroy()
            self._do_export(fmt)

        ttk.Button(dialog, text="Export", command=do_export).pack(pady=12)

    def _do_export(self, fmt: str) -> None:
        path = filedialog.asksaveasfilename(
            parent=self.window,
            title="Save export",
            defaultextension=".txt" if fmt == "Text" else ".csv" if fmt == "Anki CSV" else ".json",
            filetypes=[
                ("Text files", "*.txt") if fmt == "Text" else ("CSV files", "*.csv") if fmt == "Anki CSV" else ("JSON files", "*.json"),
                ("All files", "*.*"),
            ],
        )
        if not path:
            return

        try:
            if fmt == "Text":
                show_trans = self.show_trans_var.get()
                lines = []
                for target, trans in self.generated_pairs:
                    lines.append(target)
                    if show_trans:
                        lines.append(trans)
                    lines.append("")
                Path(path).write_text("\n".join(lines), encoding="utf-8")
            elif fmt == "Anki CSV":
                rows = ["Front,Back,Tags"]
                for target, trans in self.generated_pairs:
                    rows.append(f'"{target}","{trans}","language-learning"')
                Path(path).write_text("\n".join(rows), encoding="utf-8")
            else:  # JSON
                import json
                data = [{"front": t, "back": tr, "tags": ["language-learning"]} for t, tr in self.generated_pairs]
                Path(path).write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

            self.status_var.set(f"Exported to {Path(path).name}")
        except Exception as exc:
            messagebox.showerror("Export Error", str(exc))

    def _clear(self) -> None:
        self.output_text.delete("1.0", "end")
        self.generated_pairs.clear()
        self.status_var.set("Cleared")

    def _on_close(self) -> None:
        # Save settings
        if self.availability.available:
            self.app.settings["language_learning"] = self._collect_settings()
        save_app_settings(self.app.settings)
        if self.worker and self.worker.is_alive():
            self.stop_event.set()
            self.pause_event.set()
            self.app.player.stop(quiet=True)
            self.worker.join(timeout=1)
        # Clean up any remaining temp files
        for path in self._temp_files:
            try:
                path.unlink(missing_ok=True)
            except Exception:
                pass
        self.window.destroy()


    def _rebuild_styles(self) -> None:
        """Refresh wizard styles after theme change."""
        if hasattr(self, "output_text"):
            self.output_text.configure(
                bg=TEXT_BG,
                fg=THEMES[CURRENT_THEME]["label_fg"],
                insertbackground=ACCENT,
                highlightbackground=TEXT_BORDER,
                highlightcolor=ACCENT,
            )
        self.window.update_idletasks()


class SettingsDialog:
    def __init__(self, app: "App") -> None:
        self.app = app
        self.window = Toplevel(app.root)
        self.window.title("Settings")
        self.window.geometry("640x520")
        self.window.transient(app.root)
        self.window.grab_set()

        self.settings = app.settings.copy()

        # Notebook with tabs
        self.notebook = ttk.Notebook(self.window)
        self.notebook.pack(fill="both", expand=True, padx=12, pady=12)

        self._build_general_tab()
        self._build_audio_tab()
        self._build_appearance_tab()
        self._build_advanced_tab()

        # Button bar
        btn_frame = ttk.Frame(self.window)
        btn_frame.pack(fill="x", padx=12, pady=(0, 12))
        ttk.Button(btn_frame, text="OK", command=self._on_ok).pack(side="right")
        ttk.Button(btn_frame, text="Cancel", command=self._on_cancel).pack(side="right", padx=(0, 8))
        ttk.Button(btn_frame, text="Apply", command=self._on_apply).pack(side="right", padx=(0, 8))
        ttk.Button(btn_frame, text="Reset to Defaults", command=self._on_reset).pack(side="left")

    def _build_general_tab(self) -> None:
        frame = ttk.Frame(self.notebook, padding=12)
        self.notebook.add(frame, text="General")
        frame.columnconfigure(1, weight=1)

        row = 0
        ttk.Label(frame, text="Default Engine").grid(row=row, column=0, sticky="w", pady=6)
        self.general_engine_var = StringVar(value=self.settings.get("general", {}).get("default_engine", ENGINE_AUTO))
        ttk.Combobox(frame, textvariable=self.general_engine_var, values=[ENGINE_AUTO, ENGINE_PIPER, ENGINE_XTTS, ENGINE_POCKET], state="readonly", width=20).grid(row=row, column=1, sticky="w", pady=6)

        row += 1
        ttk.Label(frame, text="Default Language").grid(row=row, column=0, sticky="w", pady=6)
        self.general_lang_var = StringVar(value=self.settings.get("general", {}).get("default_language", "hu"))
        ttk.Combobox(frame, textvariable=self.general_lang_var, values=[language_display_name(c) for c in available_languages(self.app.piper_voice_options)], state="readonly", width=20).grid(row=row, column=1, sticky="w", pady=6)

        row += 1
        ttk.Label(frame, text="Default Output Folder").grid(row=row, column=0, sticky="w", pady=6)
        self.general_output_var = StringVar(value=self.settings.get("paths", {}).get("output_folder", str(get_default_music_folder())))
        out_frame = ttk.Frame(frame)
        out_frame.grid(row=row, column=1, sticky="ew", pady=6)
        out_frame.columnconfigure(0, weight=1)
        ttk.Entry(out_frame, textvariable=self.general_output_var).grid(row=0, column=0, sticky="ew")
        ttk.Button(out_frame, text="Browse…", command=self._pick_output_folder).grid(row=0, column=1, padx=(6, 0))

        row += 1
        self.general_autosave_var = BooleanVar(value=self.settings.get("general", {}).get("auto_save_output_folder", True))
        ttk.Checkbutton(frame, text="Auto-save output folder", variable=self.general_autosave_var).grid(row=row, column=0, columnspan=2, sticky="w", pady=6)

        row += 1
        self.general_confirm_exit_var = BooleanVar(value=self.settings.get("general", {}).get("confirm_on_exit", True))
        ttk.Checkbutton(frame, text="Confirm on exit", variable=self.general_confirm_exit_var).grid(row=row, column=0, columnspan=2, sticky="w", pady=6)

    def _pick_output_folder(self) -> None:
        folder = filedialog.askdirectory(parent=self.window, title="Choose Default Output Folder", initialdir=self.general_output_var.get())
        if folder:
            self.general_output_var.set(folder)

    def _build_audio_tab(self) -> None:
        frame = ttk.Frame(self.notebook, padding=12)
        self.notebook.add(frame, text="Audio")
        frame.columnconfigure(1, weight=1)

        row = 0
        ttk.Label(frame, text="Default Output Format").grid(row=row, column=0, sticky="w", pady=6)
        self.audio_format_var = StringVar(value=self.settings.get("audio", {}).get("default_format", "MP3"))
        ttk.Combobox(frame, textvariable=self.audio_format_var, values=list(SUPPORTED_OUTPUT_FORMATS.values()), state="readonly", width=20).grid(row=row, column=1, sticky="w", pady=6)

        row += 1
        ttk.Label(frame, text="MP3 Quality").grid(row=row, column=0, sticky="w", pady=6)
        self.audio_mp3_var = StringVar(value=self.settings.get("audio", {}).get("mp3_quality", "192 kbps (recommended)"))
        ttk.Combobox(frame, textvariable=self.audio_mp3_var, values=list(MP3_QUALITY_PRESETS.keys()), state="readonly", width=30).grid(row=row, column=1, sticky="w", pady=6)

        row += 1
        ttk.Label(frame, text="OGG Quality").grid(row=row, column=0, sticky="w", pady=6)
        self.audio_ogg_var = StringVar(value=self.settings.get("audio", {}).get("ogg_quality", "High (q5)"))
        ttk.Combobox(frame, textvariable=self.audio_ogg_var, values=list(OGG_QUALITY_PRESETS.keys()), state="readonly", width=30).grid(row=row, column=1, sticky="w", pady=6)

        row += 1
        ttk.Label(frame, text="Sample Rate").grid(row=row, column=0, sticky="w", pady=6)
        self.audio_sr_var = IntVar(value=self.settings.get("audio", {}).get("sample_rate", 44100))
        ttk.Combobox(frame, textvariable=self.audio_sr_var, values=[22050, 44100, 48000], state="readonly", width=20).grid(row=row, column=1, sticky="w", pady=6)

        row += 1
        ttk.Label(frame, text="Audio Output Device").grid(row=row, column=0, sticky="w", pady=6)
        self.audio_device_var = StringVar(value=self.settings.get("audio", {}).get("output_device", ""))
        device_frame = ttk.Frame(frame)
        device_frame.grid(row=row, column=1, sticky="ew", pady=6)
        ttk.Combobox(device_frame, textvariable=self.audio_device_var, state="readonly", width=30).grid(row=0, column=0, sticky="ew")
        ttk.Button(device_frame, text="Refresh", command=self._refresh_audio_devices).grid(row=0, column=1, padx=(6, 0))
        self._refresh_audio_devices()

    def _refresh_audio_devices(self) -> None:
        try:
            import pygame
            if not pygame.mixer.get_init():
                pygame.mixer.init()
            # Note: pygame doesn't easily expose device list in older versions
            # This is a placeholder for future enhancement
        except Exception:
            pass

    def _build_appearance_tab(self) -> None:
        frame = ttk.Frame(self.notebook, padding=12)
        self.notebook.add(frame, text="Appearance")
        frame.columnconfigure(1, weight=1)

        row = 0
        ttk.Label(frame, text="Theme").grid(row=row, column=0, sticky="w", pady=6)
        self.appearance_theme_var = StringVar(value=self.settings.get("ui", {}).get("theme", "light"))
        ttk.Combobox(frame, textvariable=self.appearance_theme_var, values=["light", "dark", "high_contrast"], state="readonly", width=20).grid(row=row, column=1, sticky="w", pady=6)

        row += 1
        ttk.Label(frame, text="Font Size").grid(row=row, column=0, sticky="w", pady=6)
        self.appearance_font_var = IntVar(value=self.settings.get("ui", {}).get("font_size", 11))
        ttk.Spinbox(frame, from_=9, to=14, textvariable=self.appearance_font_var, width=8).grid(row=row, column=1, sticky="w", pady=6)

        row += 1
        self.appearance_compact_var = BooleanVar(value=self.settings.get("ui", {}).get("toolbar_compact", False))
        ttk.Checkbutton(frame, text="Compact toolbar (icons only)", variable=self.appearance_compact_var).grid(row=row, column=0, columnspan=2, sticky="w", pady=6)

        row += 1
        self.appearance_sidebar_var = BooleanVar(value=self.settings.get("ui", {}).get("sidebar_collapsed", False))
        ttk.Checkbutton(frame, text="Sidebar collapsed by default", variable=self.appearance_sidebar_var).grid(row=row, column=0, columnspan=2, sticky="w", pady=6)

    def _build_advanced_tab(self) -> None:
        frame = ttk.Frame(self.notebook, padding=12)
        self.notebook.add(frame, text="Advanced")
        frame.columnconfigure(1, weight=1)

        row = 0
        ttk.Label(frame, text="XTTS License Accepted").grid(row=row, column=0, sticky="w", pady=6)
        self.advanced_xtts_var = BooleanVar(value=self.settings.get("general", {}).get("xtts_license_accepted", False))
        ttk.Checkbutton(frame, text="I accept the Coqui CPML license for XTTS v2", variable=self.advanced_xtts_var).grid(row=row, column=1, sticky="w", pady=6)

        row += 1
        ttk.Label(frame, text="Piper Voice Directory").grid(row=row, column=0, sticky="w", pady=6)
        self.advanced_piper_dir_var = StringVar(value=self.settings.get("paths", {}).get("piper_voice_dir", "voices/piper"))
        dir_frame = ttk.Frame(frame)
        dir_frame.grid(row=row, column=1, sticky="ew", pady=6)
        ttk.Entry(dir_frame, textvariable=self.advanced_piper_dir_var).grid(row=0, column=0, sticky="ew")
        ttk.Button(dir_frame, text="Browse…", command=lambda: self._pick_dir(self.advanced_piper_dir_var)).grid(row=0, column=1, padx=(6, 0))

        row += 1
        ttk.Label(frame, text="Pocket Voice Directory").grid(row=row, column=0, sticky="w", pady=6)
        self.advanced_pocket_dir_var = StringVar(value=self.settings.get("paths", {}).get("pocket_voice_dir", "voices/pocket"))
        dir_frame2 = ttk.Frame(frame)
        dir_frame2.grid(row=row, column=1, sticky="ew", pady=6)
        ttk.Entry(dir_frame2, textvariable=self.advanced_pocket_dir_var).grid(row=0, column=0, sticky="ew")
        ttk.Button(dir_frame2, text="Browse…", command=lambda: self._pick_dir(self.advanced_pocket_dir_var)).grid(row=0, column=1, padx=(6, 0))

        row += 1
        ttk.Label(frame, text="Log Verbosity").grid(row=row, column=0, sticky="w", pady=6)
        self.advanced_log_var = StringVar(value=self.settings.get("advanced", {}).get("log_level", "INFO"))
        ttk.Combobox(frame, textvariable=self.advanced_log_var, values=["DEBUG", "INFO", "WARNING", "ERROR"], state="readonly", width=20).grid(row=row, column=1, sticky="w", pady=6)

        row += 1
        ttk.Button(frame, text="Reset All Settings to Defaults", command=self._on_reset).grid(row=row, column=0, columnspan=2, pady=(12, 0))

    def _pick_dir(self, var: StringVar) -> None:
        folder = filedialog.askdirectory(parent=self.window, title="Choose Directory", initialdir=var.get())
        if folder:
            var.set(folder)

    def _collect_settings(self) -> dict:
        return {
            "general": {
                "default_engine": self.general_engine_var.get(),
                "default_language": language_code_from_display(self.general_lang_var.get()),
                "output_folder": self.general_output_var.get(),
                "auto_save_output_folder": self.general_autosave_var.get(),
                "confirm_on_exit": self.general_confirm_exit_var.get(),
                "xtts_license_accepted": self.advanced_xtts_var.get(),
            },
            "audio": {
                "default_format": self.audio_format_var.get(),
                "mp3_quality": self.audio_mp3_var.get(),
                "ogg_quality": self.audio_ogg_var.get(),
                "sample_rate": self.audio_sr_var.get(),
                "output_device": self.audio_device_var.get(),
            },
            "ui": {
                "theme": self.appearance_theme_var.get(),
                "font_size": self.appearance_font_var.get(),
                "toolbar_compact": self.appearance_compact_var.get(),
                "sidebar_collapsed": self.appearance_sidebar_var.get(),
            },
            "paths": {
                "piper_voice_dir": self.advanced_piper_dir_var.get(),
                "pocket_voice_dir": self.advanced_pocket_dir_var.get(),
                "output_folder": self.general_output_var.get(),
            },
            "advanced": {
                "log_level": self.advanced_log_var.get(),
            },
        }

    def _apply_settings(self, settings: dict) -> None:
        self.app.settings.update(settings)
        save_app_settings(self.app.settings)
        # Apply theme immediately if changed
        theme_name = settings.get("ui", {}).get("theme", "light")
        resolved = resolve_theme(theme_name)
        apply_theme(resolved, self.app.root)
        self.app._rebuild_styles()
        # Refresh wizard dialogs if open
        for wizard_attr in ("voice_wizard", "doc_wizard", "lang_learning_wizard"):
            wizard = getattr(self.app, wizard_attr, None)
            if wizard and wizard.window.winfo_exists():
                wizard._rebuild_styles()

    def _on_ok(self) -> None:
        self._apply_settings(self._collect_settings())
        self.window.destroy()

    def _on_cancel(self) -> None:
        self.window.destroy()

    def _on_apply(self) -> None:
        self._apply_settings(self._collect_settings())

    def _on_reset(self) -> None:
        if messagebox.askyesno("Reset Settings", "Reset all settings to defaults? This cannot be undone."):
            self.app.settings = {}
            save_app_settings(self.app.settings)
            # Reopen with defaults
            self.window.destroy()
            SettingsDialog(self.app)


class App:
    def __init__(self, root: Tk) -> None:
        self.root = root
        self.root.title("Local TTS Audio Generator")
        self.root.geometry("980x760")
        self.root.minsize(900, 650)

        self.log_queue: queue.Queue[str] = queue.Queue()
        self.service = SynthesisCoordinator(self.enqueue_log)
        self.worker: threading.Thread | None = None
        self.preview_worker: threading.Thread | None = None
        self.preview_stop_event = threading.Event()
        self.preview_job_id = 0
        self.last_selection_start_offset: int | None = None
        self.player = AudioPlayer(self.enqueue_log)
        self.settings = load_app_settings()
        self.voice_wizard: PiperVoiceWizard | None = None
        self.doc_wizard: DocumentToAudioWizard | None = None
        self.lang_learning_wizard: LanguageLearningWizard | None = None
        self.piper_voice_options = discover_local_piper_voices()

        initial_language = self.settings.get("general", {}).get("default_language", "hu")
        if initial_language not in available_languages(self.piper_voice_options):
            initial_language = "hu"
        self.language = StringVar(value=initial_language)
        self.language_display = StringVar(value=language_display_name(initial_language))
        initial_engine = self.settings.get("general", {}).get("default_engine", ENGINE_AUTO)
        self.engine = StringVar(value=initial_engine if initial_engine in ENGINE_SUMMARIES else ENGINE_AUTO)
        self._syncing_voice_settings = False
        initial_piper_voice = self.settings.get("default_piper_voice_label", DEFAULT_PIPER_VOICE_LABEL)
        if initial_piper_voice not in self.piper_voice_options:
            initial_piper_voice = DEFAULT_PIPER_VOICE_LABEL
        self.piper_voice_label = StringVar(value=initial_piper_voice)
        self.speaker_name = StringVar(value=DEFAULT_SPEAKER)
        self.pocket_voice = StringVar(value=pocket_default_voice("hu"))
        self.speaker_wav = StringVar()
        configured_output = self.settings.get("paths", {}).get("output_folder") or str(get_default_music_folder())
        self.output_file = StringVar(value=str((Path(configured_output) / "speech.mp3").resolve()))
        self.status = StringVar(value="Ready")
        self.playback_toggle_label = StringVar(value="Pause")
        self.speed = DoubleVar(value=1.0)
        self.generation_modal: Toplevel | None = None
        self.generation_progress = None
        self.generation_status = StringVar(value="")
        self.generation_result_path: Path | None = None
        self.generation_close_button = None
        self.generation_open_file_button = None
        self.generation_open_folder_button = None

        # Sidebar state
        self.sidebar_collapsed = BooleanVar(value=self.settings.get("ui", {}).get("sidebar_collapsed", False))

        # Apply theme from settings
        theme_name = self.settings.get("ui", {}).get("theme", "light")
        resolved_theme = resolve_theme(theme_name)
        apply_theme(resolved_theme, self.root)

        self._build_ui()
        self.language_display.trace_add("write", self._on_language_display_changed)
        self.engine.trace_add("write", self.on_voice_settings_changed)
        self.piper_voice_label.trace_add("write", self.on_voice_settings_changed)
        self.speaker_wav.trace_add("write", self.on_voice_settings_changed)
        self.on_voice_settings_changed()
        self._bind_shortcuts()
        self.root.after(150, self.flush_logs)

    def _on_language_display_changed(self, *_args) -> None:
        self.language.set(language_code_from_display(self.language_display.get()))
        self.on_voice_settings_changed()

    def _configure_styles(self) -> None:  # pragma: no cover
        style = ttk.Style()
        if "clam" in style.theme_names():
            style.theme_use("clam")
        apply_theme(CURRENT_THEME, self.root)

    def _rebuild_styles(self) -> None:
        """Refresh widgets that ttk cannot recolor through a style alone."""
        if hasattr(self, "text"):
            self.text.configure(
                bg=TEXT_BG,
                fg=THEMES[CURRENT_THEME]["label_fg"],
                insertbackground=ACCENT,
                highlightbackground=TEXT_BORDER,
                highlightcolor=ACCENT,
            )
            self.text.tag_configure(READ_ALOUD_LINE_TAG, background=READ_ALOUD_HIGHLIGHT)
        if hasattr(self, "log"):
            self.log.configure(
                bg=THEMES[CURRENT_THEME]["log_bg"],
                fg=THEMES[CURRENT_THEME]["log_fg"],
                insertbackground=ACCENT,
            )
        self.root.update_idletasks()

    def _bind_shortcuts(self) -> None:
        """Bind keyboard shortcuts."""
        self.root.bind_all("<Control-n>", lambda e: self.new_document())
        self.root.bind_all("<Control-o>", lambda e: self.load_text_file())
        self.root.bind_all("<Control-s>", lambda e: self.save_audio_as())
        self.root.bind_all("<Control-q>", lambda e: self.root.quit())
        self.root.bind_all("<Control-comma>", lambda e: self.open_settings())
        self.root.bind_all("<F5>", lambda e: self.start_generation())
        self.root.bind_all("<F6>", lambda e: self.start_read_aloud())
        self.root.bind_all("<Escape>", lambda e: self.stop_playback())
        self.root.bind_all("<F1>", lambda e: self.show_about())

        # Space: play/pause ONLY when text area does NOT have focus
        def on_space(event):
            if self.root.focus_get() is not self.text:
                self.toggle_playback_pause()
                return "break"
        self.root.bind_all("<space>", on_space)

    def _build_ui(self) -> None:  # pragma: no cover
        self._configure_styles()
        main = ttk.Frame(self.root)
        main.pack(fill="both", expand=True)
        main.columnconfigure(0, weight=1)
        main.rowconfigure(2, weight=1)

        # Row 0: Toolbar
        self.toolbar = self._build_toolbar(main)
        self.toolbar.grid(row=0, column=0, columnspan=2, sticky="ew")

        # Row 1: Separator
        ttk.Separator(main, orient="horizontal").grid(row=1, column=0, columnspan=2, sticky="ew", pady=(0, 4))

        # Row 2: Text area (left) + Sidebar (right)
        self.text_area_frame = ttk.Frame(main)
        self.text_area_frame.grid(row=2, column=0, sticky="nsew")
        self.text_area_frame.columnconfigure(0, weight=1)
        self.text_area_frame.rowconfigure(0, weight=1)

        self.sidebar_frame = ttk.Frame(main, style="Sidebar.TFrame", width=320)
        self.sidebar_frame.grid(row=2, column=1, sticky="ns")
        self.sidebar_frame.grid_propagate(False)
        self.sidebar_frame.columnconfigure(0, weight=1)

        # Build sidebar content (Options)
        self._build_sidebar()

        # Text area
        self._build_text_area()

        # Row 3: Status/Log panel
        log_frame = ttk.LabelFrame(main, text="Status", padding=10)
        log_frame.grid(row=3, column=0, columnspan=2, sticky="ew", pady=(8, 0))
        log_frame.columnconfigure(0, weight=1)
        ttk.Label(log_frame, textvariable=self.status).grid(row=0, column=0, sticky="w")
        self.log = Text(
            log_frame,
            height=10,
            wrap="word",
            font=(FONT_MONO, 10),
            bg=THEMES[CURRENT_THEME]["log_bg"],
            fg=THEMES[CURRENT_THEME]["log_fg"],
            relief="flat",
            insertbackground=ACCENT,
        )
        self.log.grid(row=1, column=0, sticky="ew", pady=(8, 0))
        self.log.configure(state="disabled")

        # Apply initial sidebar state
        if self.sidebar_collapsed.get():
            self._collapse_sidebar()
        else:
            self._expand_sidebar()

    def _build_toolbar(self, parent) -> ttk.Frame:
        toolbar = ttk.Frame(parent, style="Toolbar.TFrame", padding=(8, 6))
        toolbar.columnconfigure(1, weight=1)  # spacer

        # File menubutton (Start Menu)
        file_menu = Menu(toolbar, tearoff=0)
        file_menu.add_command(label="New\t\tCtrl+N", command=self.new_document)
        file_menu.add_command(label="Open Text…\tCtrl+O", command=self.load_text_file)
        file_menu.add_command(label="Save Audio As…\tCtrl+S", command=self.save_audio_as)
        file_menu.add_separator()
        file_menu.add_command(label="Export Log…", command=self.export_log)
        file_menu.add_separator()
        file_menu.add_command(label="Settings…\tCtrl+,", command=self.open_settings)
        file_menu.add_separator()
        file_menu.add_command(label="Voice Wizard…", command=self.open_voice_wizard)
        file_menu.add_command(label="Language Learning…", command=self.open_language_learning)
        file_menu.add_command(label="Document Converter…", command=self.open_document_wizard)
        file_menu.add_separator()
        file_menu.add_command(label="Exit\t\tCtrl+Q", command=self._confirm_exit)
        file_mb = ttk.Menubutton(toolbar, text="File", menu=file_menu, direction="below")
        file_mb.grid(row=0, column=0, padx=(0, 8))

        # Home actions
        home_frame = ttk.Frame(toolbar)
        home_frame.grid(row=0, column=1, sticky="w")
        ttk.Button(home_frame, text="Load text", command=self.load_text_file).pack(side="left", padx=2)
        ttk.Button(home_frame, text="Generate audio", style="Accent.TButton", command=self.start_generation).pack(side="left", padx=2)
        ttk.Button(home_frame, text="Read aloud", command=self.start_read_aloud).pack(side="left", padx=2)
        self.playback_button = ttk.Button(home_frame, textvariable=self.playback_toggle_label, command=self.toggle_playback_pause)
        self.playback_button.pack(side="left", padx=2)
        ttk.Button(home_frame, text="Stop", command=self.stop_playback).pack(side="left", padx=2)

        # Voice menubutton
        voice_menu = Menu(toolbar, tearoff=0)
        engine_submenu = Menu(voice_menu, tearoff=0)
        for eng in [ENGINE_AUTO, ENGINE_PIPER, ENGINE_XTTS, ENGINE_POCKET]:
            engine_submenu.add_radiobutton(label=eng, variable=self.engine, value=eng, command=self.on_voice_settings_changed)
        voice_menu.add_cascade(label="Engine", menu=engine_submenu)
        voice_menu.add_separator()
        voice_menu.add_command(label="Voice Wizard…", command=self.open_voice_wizard)
        voice_menu.add_command(label="Piper Voice Manager…", command=self.open_voice_wizard)
        voice_mb = ttk.Menubutton(toolbar, text="Voice", menu=voice_menu, direction="below")
        voice_mb.grid(row=0, column=2, padx=(8, 4))

        # Tools menubutton
        tools_menu = Menu(toolbar, tearoff=0)
        tools_menu.add_command(label="Language Learning…", command=self.open_language_learning)
        tools_menu.add_command(label="Document Converter…", command=self.open_document_wizard)
        tools_mb = ttk.Menubutton(toolbar, text="Tools", menu=tools_menu, direction="below")
        tools_mb.grid(row=0, column=3, padx=4)

        # Settings button
        ttk.Button(toolbar, text="Settings", command=self.open_settings).grid(row=0, column=4, padx=(8, 4))

        # Sidebar toggle button
        self.sidebar_toggle_btn = ttk.Button(toolbar, text="◀", width=3, command=self._toggle_sidebar)
        self.sidebar_toggle_btn.grid(row=0, column=5, padx=(4, 0))

        return toolbar

    def _build_sidebar(self) -> None:
        """Build the collapsible sidebar with Options controls."""
        sidebar = self.sidebar_frame
        sidebar.columnconfigure(0, weight=1)

        # Title
        ttk.Label(sidebar, text="Options", style="Header.TLabel").grid(row=0, column=0, sticky="w", padx=12, pady=(12, 8))

        controls = ttk.Frame(sidebar, padding=(12, 0, 12, 12))
        controls.grid(row=1, column=0, sticky="nsew")
        controls.columnconfigure(1, weight=1)

        # Language
        ttk.Label(controls, text="Language").grid(row=0, column=0, sticky="w", pady=6)
        self.lang_box = ttk.Combobox(
            controls,
            textvariable=self.language_display,
            values=[language_display_name(c) for c in available_languages(self.piper_voice_options)],
            state="readonly",
            width=16,
        )
        self.lang_box.grid(row=0, column=1, sticky="ew", pady=6)

        # Engine
        ttk.Label(controls, text="Engine").grid(row=1, column=0, sticky="w", pady=6)
        self.engine_box = ttk.Combobox(
            controls,
            textvariable=self.engine,
            values=[ENGINE_AUTO, ENGINE_PIPER, ENGINE_XTTS, ENGINE_POCKET],
            state="readonly",
            width=18,
        )
        self.engine_box.grid(row=1, column=1, sticky="ew", pady=6)

        # Voice (adaptive)
        self.voice_label_var = StringVar(value="Voice")
        ttk.Label(controls, textvariable=self.voice_label_var).grid(row=2, column=0, sticky="w", pady=6)
        voice_area = ttk.Frame(controls)
        voice_area.grid(row=2, column=1, sticky="ew", pady=6)
        voice_area.columnconfigure(0, weight=1)

        self.piper_voice_box = ttk.Combobox(
            voice_area,
            textvariable=self.piper_voice_label,
            values=list(self.piper_voice_options.keys()),
            state="readonly",
        )
        self.pocket_voice_box = ttk.Combobox(
            voice_area,
            textvariable=self.pocket_voice,
            values=POCKET_PREDEFINED_VOICES,
            state="readonly",
        )
        self.speaker_name_entry = ttk.Entry(voice_area, textvariable=self.speaker_name)
        for widget in (self.piper_voice_box, self.pocket_voice_box, self.speaker_name_entry):
            widget.grid(row=0, column=0, sticky="ew")
            widget.grid_remove()
        add_context_menu(self.speaker_name_entry)

        # Reference WAV
        ttk.Label(controls, text="Reference WAV").grid(row=3, column=0, sticky="w", pady=6)
        self.speaker_wav_entry = ttk.Entry(controls, textvariable=self.speaker_wav)
        self.speaker_wav_entry.grid(row=3, column=1, sticky="ew", pady=6)
        add_context_menu(self.speaker_wav_entry)
        self.speaker_wav_button = ttk.Button(controls, text="Browse", command=self.pick_reference_wav)
        self.speaker_wav_button.grid(row=3, column=2, sticky="e", pady=6, padx=(4, 0))

        # Output file
        ttk.Label(controls, text="Output file").grid(row=4, column=0, sticky="w", pady=6)
        output_file_entry = ttk.Entry(controls, textvariable=self.output_file)
        output_file_entry.grid(row=4, column=1, sticky="ew", pady=6)
        add_context_menu(output_file_entry)
        ttk.Button(controls, text="Save As", command=self.pick_output_file).grid(row=4, column=2, sticky="e", pady=6, padx=(4, 0))

        # Speed
        self.speed_label = ttk.Label(controls, text="Speed: 1.0x")
        self.speed_label.grid(row=5, column=0, sticky="w", pady=6)
        self.speed_slider = ttk.Scale(
            controls,
            from_=0.5,
            to=2.0,
            variable=self.speed,
            orient="horizontal",
            command=self._on_speed_changed,
        )
        self.speed_slider.grid(row=5, column=1, columnspan=2, sticky="ew", pady=6)

        # Engine hint
        self.engine_hint = StringVar(value="")
        ttk.Label(controls, textvariable=self.engine_hint, style="Hint.TLabel").grid(
            row=6, column=0, columnspan=3, sticky="w", pady=(8, 0)
        )

    def _build_text_area(self) -> None:
        """Build the main text editing area."""
        ttk.Label(self.text_area_frame, text="Text").grid(row=0, column=0, sticky="w", padx=8, pady=(8, 0))
        self.textbox = ttk.Frame(self.text_area_frame)
        self.textbox.grid(row=1, column=0, sticky="nsew", padx=8, pady=(0, 8))
        self.textbox.columnconfigure(0, weight=1)
        self.textbox.rowconfigure(0, weight=1)

        self.text = Text(
            self.textbox,
            wrap="word",
            font=(FONT_BODY, 11),
            padx=12,
            pady=12,
            undo=True,
            exportselection=False,
            bg=TEXT_BG,
            fg=THEMES[CURRENT_THEME]["label_fg"],
            relief="flat",
            insertbackground=ACCENT,
            highlightthickness=1,
            highlightbackground=TEXT_BORDER,
            highlightcolor=ACCENT,
        )
        self.text.grid(row=0, column=0, sticky="nsew")
        add_context_menu(self.text)
        text_scroll = ttk.Scrollbar(self.textbox, orient="vertical", command=self.text.yview)
        text_scroll.grid(row=0, column=1, sticky="ns")
        self.text.configure(yscrollcommand=text_scroll.set)
        self.text.tag_configure(READ_ALOUD_LINE_TAG, background=READ_ALOUD_HIGHLIGHT)
        self.text.bind("<ButtonRelease-1>", self.on_text_click)
        self.text.bind("<ButtonRelease-1>", self.update_selection_cache, add="+")
        self.text.bind("<KeyRelease>", self.update_selection_cache, add="+")

    def _toggle_sidebar(self) -> None:
        if self.sidebar_collapsed.get():
            self._expand_sidebar()
        else:
            self._collapse_sidebar()

    def _collapse_sidebar(self) -> None:
        self.sidebar_frame.grid_remove()
        self.sidebar_toggle_btn.configure(text="▶")
        self.sidebar_collapsed.set(True)
        self.settings.setdefault("ui", {})["sidebar_collapsed"] = True
        save_app_settings(self.settings)

    def _expand_sidebar(self) -> None:
        self.sidebar_frame.grid()
        self.sidebar_toggle_btn.configure(text="◀")
        self.sidebar_collapsed.set(False)
        self.settings.setdefault("ui", {})["sidebar_collapsed"] = False
        save_app_settings(self.settings)

    def _confirm_exit(self) -> None:
        """Confirm exit if setting is enabled."""
        if self.settings.get("general", {}).get("confirm_on_exit", True):
            if not messagebox.askyesno("Exit", "Are you sure you want to exit?"):
                return
        self.root.quit()

    def new_document(self) -> None:
        """Clear the text area for a new document."""
        self.text.delete("1.0", "end")
        self.last_selection_start_offset = None
        self.clear_read_aloud_highlight()
        self.enqueue_log("New document created.")

    def save_audio_as(self) -> None:
        """Pick output file and generate audio."""
        self.pick_output_file()
        self.start_generation()

    def export_log(self) -> None:
        """Export the status log to a file."""
        path = filedialog.asksaveasfilename(
            parent=self.root,
            title="Export Log",
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
        )
        if not path:
            return
        try:
            log_content = self.log.get("1.0", "end-1c")
            Path(path).write_text(log_content, encoding="utf-8")
            self.enqueue_log(f"Log exported to {path}")
        except Exception as exc:
            messagebox.showerror("Export Failed", str(exc))

    def open_settings(self) -> None:
        """Open the settings dialog."""
        SettingsDialog(self)

    def show_about(self) -> None:
        """Show the about dialog."""
        messagebox.showinfo(
            "About Local TTS Audio Generator",
            "Local TTS Audio Generator\n\n"
            "A desktop text-to-speech application supporting:\n"
            "• Piper TTS (fast, offline)\n"
            "• XTTS v2 (high quality, voice cloning)\n"
            "• Pocket TTS (lightweight neural voices)\n\n"
            "Built with Python, tkinter, and Coqui TTS.",
        )

    def resolved_engine(self) -> str:
        return select_engine(
            self.engine.get(),
            self.language.get(),
            bool(self.speaker_wav.get().strip()),
            self.piper_voice_options,
        )

    def reload_piper_voices(self, preferred_code: str | None = None) -> None:
        self.piper_voice_options = discover_local_piper_voices()
        labels = list(self.piper_voice_options.keys())
        self.piper_voice_box.configure(values=labels)
        if hasattr(self, "lang_box"):
            self.lang_box.configure(
                values=[language_display_name(c) for c in available_languages(self.piper_voice_options)]
            )

        preferred_label = self.piper_voice_label.get()
        if preferred_code is not None:
            preferred_label = self.find_piper_label_by_code(preferred_code) or preferred_label
        if preferred_label not in self.piper_voice_options:
            preferred_label = DEFAULT_PIPER_VOICE_LABEL
        # A freshly downloaded voice may be in a new language — follow it so the
        # language-first UI stays coherent.
        if (
            preferred_code is not None
            and preferred_label in self.piper_voice_options
            and hasattr(self, "language_display")
        ):
            code = self.piper_voice_options[preferred_label]["code"]
            self.language_display.set(language_display_name(piper_language_of_code(code)))
        self.piper_voice_label.set(preferred_label)

    def find_piper_label_by_code(self, voice_code: str) -> str | None:
        for label, metadata in self.piper_voice_options.items():
            if metadata["code"] == voice_code:
                return label
        return None

    def set_default_piper_voice(self, label: str) -> None:
        self.settings["default_piper_voice_label"] = label
        save_app_settings(self.settings)
        self.piper_voice_label.set(label)

    def on_voice_settings_changed(self, *_args) -> None:
        if getattr(self, "_syncing_voice_settings", False):
            return
        self._syncing_voice_settings = True
        try:
            self._sync_voice_settings()
        finally:
            self._syncing_voice_settings = False

    def _sync_voice_settings(self) -> None:
        language = self.language.get()

        # 1. Offer only engines that can actually speak the chosen language.
        supported = engines_supporting_language(language, self.piper_voice_options)
        engine_values = [ENGINE_AUTO] + supported
        if hasattr(self, "engine_box"):
            self.engine_box.configure(values=engine_values)
        if self.engine.get() not in engine_values:
            self.enqueue_log(
                f"{self.engine.get()} can't speak {language_display_name(language)} "
                "— switched to Auto."
            )
            self.engine.set(ENGINE_AUTO)

        resolved = self.resolved_engine()

        # 2. Piper voice list is scoped to the language; keep a valid selection.
        if resolved == ENGINE_PIPER:
            labels = piper_voices_for_language(self.piper_voice_options, language) or list(
                self.piper_voice_options.keys()
            )
            self.piper_voice_box.configure(values=labels)
            if self.piper_voice_label.get() not in labels and labels:
                self.piper_voice_label.set(labels[0])

        # 3. Point Pocket at the language's default voice, unless the user has
        #    deliberately chosen a non-default one.
        if resolved == ENGINE_POCKET and is_pocket_default_voice(self.pocket_voice.get().strip()):
            self.pocket_voice.set(pocket_default_voice(language))

        # 4. Reveal the one voice control the engine uses; relabel and toggle it.
        self._show_voice_widget(resolved)

        # 5. A reference clip clones a voice (XTTS/Pocket). Keep it enabled in
        #    Auto too, since dropping in a clip is how Auto switches to XTTS.
        clone_applicable = self.engine.get() == ENGINE_AUTO or resolved in (ENGINE_XTTS, ENGINE_POCKET)
        clone_state = "normal" if clone_applicable else "disabled"
        self.speaker_wav_entry.configure(state=clone_state)
        self.speaker_wav_button.configure(state=clone_state)

        # 6. Tell the user what the resolved engine can do.
        summary = ENGINE_SUMMARIES.get(resolved, "")
        if self.engine.get() == ENGINE_AUTO and resolved in ENGINE_SUMMARIES:
            summary = f"Auto → {resolved}. {summary}"
        self.engine_hint.set(summary)

    def _show_voice_widget(self, resolved: str) -> None:
        # (widget, label, active-state) per engine. The active control is shown
        # and enabled; the others are hidden and disabled.
        specs = {
            ENGINE_PIPER: (getattr(self, "piper_voice_box", None), "Piper voice", "readonly"),
            ENGINE_POCKET: (getattr(self, "pocket_voice_box", None), "Pocket voice", "readonly"),
            ENGINE_XTTS: (getattr(self, "speaker_name_entry", None), "Built-in speaker", "normal"),
        }
        active = specs.get(resolved, specs[ENGINE_PIPER])
        if hasattr(self, "voice_label_var"):
            self.voice_label_var.set(active[1])
        for engine_key, (widget, _label, active_state) in specs.items():
            if widget is None:
                continue
            if engine_key == resolved:
                widget.configure(state=active_state)
                widget.grid()
            else:
                widget.configure(state="disabled")
                widget.grid_remove()

    def _on_speed_changed(self, *args) -> None:
        speed_val = round(self.speed.get(), 1)
        self.speed_label.configure(text=f"Speed: {speed_val:.1f}x")

    def enqueue_log(self, message: str) -> None:
        self.log_queue.put(message)

    def flush_logs(self) -> None:
        while not self.log_queue.empty():
            line = self.log_queue.get_nowait()
            self.status.set(line)
            self.log.configure(state="normal")
            self.log.insert(END, f"{line}\n")
            self.log.see(END)
            self.log.configure(state="disabled")
        self.root.after(150, self.flush_logs)

    def pick_reference_wav(self) -> None:
        path = filedialog.askopenfilename(
            parent=self.root,
            title="Choose reference voice WAV",
            filetypes=[("Audio files", "*.wav *.mp3 *.m4a *.flac"), ("All files", "*.*")],
        )
        if path:
            self.speaker_wav.set(path)

    def pick_output_file(self) -> None:
        path = filedialog.asksaveasfilename(
            parent=self.root,
            title="Save output audio",
            defaultextension=".mp3",
            filetypes=[
                ("Audio files", "*.mp3 *.ogg *.wav"),
                ("MP3 files", "*.mp3"),
                ("OGG files", "*.ogg"),
                ("WAV files", "*.wav"),
            ],
            initialfile="speech.mp3",
        )
        if path:
            self.output_file.set(path)

    def open_voice_wizard(self) -> None:
        if self.voice_wizard is not None and self.voice_wizard.window.winfo_exists():
            self.voice_wizard.window.lift()
            self.voice_wizard.window.focus_force()
            return
        self.voice_wizard = PiperVoiceWizard(self)

    def open_document_wizard(self) -> None:
        if self.doc_wizard is not None and self.doc_wizard.window.winfo_exists():
            self.doc_wizard.window.lift()
            self.doc_wizard.window.focus_force()
            return
        self.doc_wizard = DocumentToAudioWizard(self)

    def open_language_learning(self) -> None:
        if self.lang_learning_wizard is not None and self.lang_learning_wizard.window.winfo_exists():
            self.lang_learning_wizard.window.lift()
            return
        self.lang_learning_wizard = LanguageLearningWizard(self)

    def open_path_in_system(self, path: Path) -> None:  # pragma: no cover
        if sys.platform.startswith("win"):
            os.startfile(str(path))
            return
        if sys.platform == "darwin":
            subprocess.Popen(["open", str(path)])
            return
        subprocess.Popen(["xdg-open", str(path)])

    def show_generation_modal(self, output_path: Path) -> None:  # pragma: no cover
        if self.generation_modal is not None and self.generation_modal.winfo_exists():
            self.generation_modal.destroy()

        self.generation_result_path = output_path
        self.generation_modal = Toplevel(self.root)
        self.generation_modal.title("Generating Audio")
        self.generation_modal.transient(self.root)
        self.generation_modal.geometry("520x220")
        self.generation_modal.resizable(False, False)
        self.generation_modal.grab_set()
        self.generation_modal.protocol("WM_DELETE_WINDOW", lambda: None)

        frame = ttk.Frame(self.generation_modal, padding=16)
        frame.pack(fill="both", expand=True)

        ttk.Label(frame, text="Creating audio file", style="Header.TLabel").pack(anchor="w")
        ttk.Label(
            frame,
            text=f"Output: {output_path}",
            style="Hint.TLabel",
            wraplength=480,
        ).pack(anchor="w", pady=(6, 10))

        self.generation_status.set("Preparing synthesis...")
        ttk.Label(frame, textvariable=self.generation_status).pack(anchor="w")

        self.generation_progress = ttk.Progressbar(frame, orient="horizontal", mode="determinate", maximum=100)
        self.generation_progress.pack(fill="x", pady=(10, 12))
        self.generation_progress["value"] = 0

        buttons = ttk.Frame(frame)
        buttons.pack(fill="x", side="bottom", pady=(8, 0))
        self.generation_close_button = ttk.Button(buttons, text="Close", command=self.close_generation_modal, state="disabled")
        self.generation_close_button.pack(side="right")
        self.generation_open_folder_button = ttk.Button(
            buttons,
            text="Open Folder",
            command=self.open_generated_folder,
            state="disabled",
        )
        self.generation_open_folder_button.pack(side="right", padx=(0, 8))
        self.generation_open_file_button = ttk.Button(
            buttons,
            text="Open File",
            command=self.open_generated_file,
            state="disabled",
        )
        self.generation_open_file_button.pack(side="right", padx=(0, 8))

        self.generation_modal.update_idletasks()
        x = self.root.winfo_rootx() + max(0, (self.root.winfo_width() - self.generation_modal.winfo_width()) // 2)
        y = self.root.winfo_rooty() + max(0, (self.root.winfo_height() - self.generation_modal.winfo_height()) // 2)
        self.generation_modal.geometry(f"+{x}+{y}")

    def close_generation_modal(self) -> None:
        if self.generation_modal is not None and self.generation_modal.winfo_exists():
            self.generation_modal.grab_release()
            self.generation_modal.destroy()
        self.generation_modal = None

    def update_generation_progress(self, current: int, total: int, message: str) -> None:
        if self.generation_modal is None or not self.generation_modal.winfo_exists() or self.generation_progress is None:
            return

        progress = 10 if total <= 0 else min(95, max(10, round((current / total) * 90)))
        self.generation_progress["value"] = progress
        self.generation_status.set(message)

    def finish_generation_modal(self, result: Path | None, error: str | None = None) -> None:
        if self.generation_modal is None or not self.generation_modal.winfo_exists() or self.generation_progress is None:
            return

        if error is None and result is not None:
            self.generation_progress["value"] = 100
            self.generation_status.set(f"Audio created: {result}")
            self.generation_result_path = result
            self.generation_modal.protocol("WM_DELETE_WINDOW", self.close_generation_modal)
            self.generation_close_button.configure(state="normal")
            self.generation_open_file_button.configure(state="normal")
            self.generation_open_folder_button.configure(state="normal")
        else:
            self.generation_progress["value"] = 0
            self.generation_status.set(f"Generation failed: {error}")
            self.generation_result_path = None
            self.generation_modal.protocol("WM_DELETE_WINDOW", self.close_generation_modal)
            self.generation_close_button.configure(state="normal")
            self.generation_open_file_button.configure(state="disabled")
            self.generation_open_folder_button.configure(state="disabled")

    def open_generated_file(self) -> None:
        if self.generation_result_path is None:
            return
        self.open_path_in_system(self.generation_result_path)

    def open_generated_folder(self) -> None:
        target = self.generation_result_path.parent if self.generation_result_path is not None else Path.cwd()
        self.open_path_in_system(target)

    def load_text_file(self) -> None:
        path = filedialog.askopenfilename(
            parent=self.root,
            title="Open text file",
            filetypes=[("Text files", "*.txt *.md"), ("All files", "*.*")],
        )
        if not path:
            return
        content = Path(path).read_text(encoding="utf-8")
        self.text.delete("1.0", END)
        self.text.insert("1.0", content)
        self.last_selection_start_offset = None
        self.clear_read_aloud_highlight()
        self.enqueue_log(f"Loaded text from {path}")

    def collect_request(self, require_output: bool = True) -> SynthesisRequest | None:
        text = self.get_text_content()
        if not text.strip():
            messagebox.showerror("Missing text", "Paste or load some text first.")
            return None

        output = self.output_file.get().strip()
        if require_output:
            if not output:
                messagebox.showerror("Missing output", "Choose an output audio file.")
                return None
            try:
                output_format_for_path(Path(output))
            except ValueError as exc:
                messagebox.showerror("Invalid output format", str(exc))
                return None

        speaker_wav = self.speaker_wav.get().strip()
        if speaker_wav and not Path(speaker_wav).exists():
            messagebox.showerror("Missing file", "The selected reference voice file does not exist.")
            return None

        # Pocket draws its built-in voice from its own dropdown; XTTS uses the
        # speaker-name entry. Feed whichever fits the resolved engine.
        if self.resolved_engine() == ENGINE_POCKET:
            speaker_name = getattr(self, "pocket_voice", self.speaker_name).get().strip() or pocket_default_voice(
                self.language.get()
            )
        else:
            speaker_name = self.speaker_name.get().strip() or DEFAULT_SPEAKER

        return SynthesisRequest(
            text=text,
            language=self.language.get().strip() or "hu",
            output_file=Path(output or PREVIEW_OUTPUT_PATH),
            engine=self.engine.get().strip() or ENGINE_AUTO,
            piper_voice_label=self.piper_voice_label.get().strip() or DEFAULT_PIPER_VOICE_LABEL,
            piper_voice_code=(self.piper_voice_options.get(
                self.piper_voice_label.get().strip() or DEFAULT_PIPER_VOICE_LABEL,
                get_piper_voice_metadata(DEFAULT_PIPER_VOICE_LABEL),
            )["code"]),
            speaker_name=speaker_name,
            speaker_wav=speaker_wav,
            speed=round(self.speed.get(), 1),
        )

    def start_generation(self) -> None:
        if self.worker and self.worker.is_alive():
            messagebox.showinfo("Busy", "Generation is already running.")
            return

        request = self.collect_request(require_output=True)
        if request is None:
            return

        if self.service.resolve_engine(request) == ENGINE_XTTS and not self.ensure_xtts_license_acceptance():
            return

        self.enqueue_log("Starting synthesis job.")
        self.show_generation_modal(request.output_file)
        self.worker = threading.Thread(target=self.run_generation, args=(request,), daemon=True)
        self.worker.start()

    def start_read_aloud(self) -> None:
        self.start_read_aloud_from(None, reason="button")

    def start_read_aloud_from(self, start_offset: int | None, reason: str = "button") -> None:
        if self.worker and self.worker.is_alive():
            messagebox.showinfo("Busy", "Audio generation is already running.")
            return

        request = self.collect_request(require_output=False)
        if request is None:
            return

        resolved_start_offset = start_offset
        if resolved_start_offset is None:
            resolved_start_offset = self.get_read_aloud_start_offset()
        if resolved_start_offset is None:
            messagebox.showinfo("No speech content", "Enter some text first.")
            return

        if self.service.resolve_engine(request) == ENGINE_XTTS and not self.ensure_xtts_license_acceptance():
            return

        self.preview_job_id += 1
        self.preview_stop_event.set()
        self.player.stop(quiet=True)
        self.cleanup_preview_files()
        self.preview_stop_event = threading.Event()
        job_id = self.preview_job_id
        if reason == "click_jump":
            self.enqueue_log("Jumping read aloud to clicked word.")
        elif self.text.tag_ranges("sel") or self.last_selection_start_offset is not None:
            self.enqueue_log("Preparing read aloud preview from selection.")
        else:
            self.enqueue_log("Preparing read aloud preview from the beginning.")
        self.preview_worker = threading.Thread(
            target=self.run_read_aloud,
            args=(request, resolved_start_offset, job_id, self.preview_stop_event),
            daemon=True,
        )
        self.preview_worker.start()
        self.update_playback_toggle_label()

    def run_read_aloud(
        self,
        request: SynthesisRequest,
        start_offset: int,
        job_id: int,
        stop_event: threading.Event,
    ) -> None:
        preview_paths: list[Path] = []
        try:
            preview_dir = PREVIEW_OUTPUT_PATH.resolve().parent
            for index, (chunk, segment) in enumerate(self.service.iter_segments(request, start_offset=start_offset), start=1):
                if stop_event.is_set() or job_id != self.preview_job_id:
                    break

                self.root.after(0, lambda offset=chunk.start: self.highlight_read_aloud_line(offset))
                preview_path = preview_dir / f"read-aloud-preview-{job_id}-{index}.wav"
                preview_paths.append(preview_path)
                export_audio_segment(segment, preview_path)
                self.root.after(0, self.update_playback_toggle_label)
                self.player.play_blocking(preview_path, stop_event)
        except Exception as exc:
            self.enqueue_log(f"Error: {exc}")
            error_message = str(exc)
            self.root.after(0, lambda message=error_message: messagebox.showerror("Read aloud failed", message))
        finally:
            if preview_paths:
                self.root.after(750, lambda paths=tuple(preview_paths): self.cleanup_preview_files(paths))
            if job_id == self.preview_job_id:
                self.root.after(0, self.clear_read_aloud_highlight)
            self.root.after(0, self.update_playback_toggle_label)

    def pause_playback(self) -> None:
        try:
            self.player.pause()
            self.update_playback_toggle_label()
        except Exception as exc:
            self.enqueue_log(f"Error: {exc}")

    def resume_playback(self) -> None:
        try:
            self.player.resume()
            self.update_playback_toggle_label()
        except Exception as exc:
            self.enqueue_log(f"Error: {exc}")

    def stop_playback(self) -> None:
        try:
            self.preview_stop_event.set()
            self.player.stop()
            self.clear_read_aloud_highlight()
            self.update_playback_toggle_label()
        except Exception as exc:
            self.enqueue_log(f"Error: {exc}")

    def toggle_playback_pause(self) -> None:
        if self.player.is_paused():
            self.resume_playback()
            return
        self.pause_playback()

    def update_playback_toggle_label(self) -> None:
        if self.player.is_paused():
            self.playback_toggle_label.set("Resume")
        else:
            self.playback_toggle_label.set("Pause")

    def get_text_content(self) -> str:
        return self.text.get("1.0", "end-1c")

    def text_index_to_offset(self, index: str) -> int:
        return int(self.text.count("1.0", index, "chars")[0])

    def offset_to_text_index(self, offset: int) -> str:
        return f"1.0 + {offset} chars"

    def get_read_aloud_start_offset(self, widget_index: str | None = None) -> int | None:
        content = self.get_text_content()
        if not content.strip():
            return None

        if widget_index is None:
            if self.text.tag_ranges("sel"):
                index = self.text.index("sel.first")
            elif self.last_selection_start_offset is not None:
                return find_word_start_offset(content, self.last_selection_start_offset)
            else:
                return find_word_start_offset(content, 0)
        else:
            index = widget_index
        offset = self.text_index_to_offset(index)
        return find_word_start_offset(content, offset)

    def update_selection_cache(self, _event=None) -> None:
        if self.preview_worker and self.preview_worker.is_alive():
            return
        if self.text.tag_ranges("sel"):
            self.last_selection_start_offset = self.text_index_to_offset("sel.first")
        else:
            self.last_selection_start_offset = None

    def cleanup_preview_files(self, paths: tuple[Path, ...] | list[Path] | None = None, retries: int = 5) -> None:
        candidates = list(paths) if paths is not None else list(PREVIEW_OUTPUT_PATH.resolve().parent.glob(PREVIEW_FILE_GLOB))
        if not candidates:
            return

        locked: list[Path] = []
        for path in candidates:
            try:
                path.unlink(missing_ok=True)
            except PermissionError:
                locked.append(path)

        if locked and retries > 0:
            self.root.after(400, lambda pending=tuple(locked), remaining=retries - 1: self.cleanup_preview_files(pending, remaining))

    def clear_read_aloud_highlight(self) -> None:
        self.text.tag_remove(READ_ALOUD_LINE_TAG, "1.0", END)

    def highlight_read_aloud_line(self, offset: int) -> None:
        target_index = self.offset_to_text_index(offset)
        line_start = self.text.index(f"{target_index} linestart")
        line_end = self.text.index(f"{target_index} lineend +1c")
        self.clear_read_aloud_highlight()
        self.text.tag_add(READ_ALOUD_LINE_TAG, line_start, line_end)
        self.text.mark_set("insert", target_index)
        self.text.see(target_index)

    def on_text_click(self, event) -> None:
        index = self.text.index(f"@{event.x},{event.y}")
        start_offset = self.get_read_aloud_start_offset(index)
        if start_offset is None:
            return

        if self.preview_worker and self.preview_worker.is_alive():
            self.highlight_read_aloud_line(start_offset)
            self.root.after(0, lambda offset=start_offset: self.start_read_aloud_from(offset, reason="click_jump"))

    def ensure_xtts_license_acceptance(self) -> bool:
        if os.environ.get("COQUI_TOS_AGREED") == "1":
            return True

        accepted = messagebox.askyesno(
            "XTTS License Confirmation",
            (
                "XTTS v2 requires you to confirm Coqui's license terms on first download.\n\n"
                "Select Yes only if either:\n"
                "- you purchased a commercial license from Coqui, or\n"
                "- you agree to the non-commercial CPML terms at https://coqui.ai/cpml\n\n"
                "Continue?"
            ),
        )
        if accepted:
            os.environ["COQUI_TOS_AGREED"] = "1"
            return True
        self.enqueue_log("XTTS license confirmation was declined.")
        return False

    def run_generation(self, request: SynthesisRequest) -> None:
        try:
            chunks = chunk_text_with_offsets(request.text)
            total_chunks = len(chunks)
            if total_chunks == 0:
                raise ValueError("Text is empty after cleanup.")

            combined = AudioSegment.silent(duration=0)
            for index, (_chunk, segment) in enumerate(self.service.iter_segments(request), start=1):
                combined += segment
                if index < total_chunks:
                    combined += AudioSegment.silent(duration=PAUSE_MS)
                self.root.after(
                    0,
                    lambda current=index, total=total_chunks: self.update_generation_progress(
                        current,
                        total,
                        f"Synthesizing chunk {current}/{total}...",
                    ),
                )

            self.root.after(0, lambda: self.generation_status.set("Saving audio file..."))
            self.enqueue_log(f"Exporting audio to {request.output_file}")
            export_audio_segment(combined, request.output_file)
            self.enqueue_log("Finished.")
            result = request.output_file
        except Exception as exc:
            self.enqueue_log(f"Error: {exc}")
            error_message = str(exc)
            self.root.after(0, lambda message=error_message: self.finish_generation_modal(None, error=message))
            return

        self.enqueue_log(f"Saved audio: {result}")
        self.root.after(0, lambda path=result: self.finish_generation_modal(path))


def main() -> None:
    root = Tk()
    App(root)
    root.mainloop()


if __name__ == "__main__":  # pragma: no cover
    main()
